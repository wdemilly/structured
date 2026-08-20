from __future__ import annotations

import difflib
import io
import json
import os
import re
import sqlite3
import statistics
import time
import uuid
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import pandas as pd
import streamlit as st
from docx import Document

try:
    from pangram import Pangram
except Exception:
    Pangram = None

try:
    from anthropic import Anthropic
except Exception:
    Anthropic = None


APP_TITLE = "Pangram Experiment Lab"
APP_VERSION = "v2.11 · web-parity harness experiment"
MIN_PANGRAM_WORDS = 50
DB_PATH = Path(__file__).with_name("pangram_microscope.db")
DEFAULT_SAMPLE_SIZES = [150]
ALL_SAMPLE_SIZES = [50, 75, 100, 125, 150, 200, 250, 300, 400, 500, 750, 1000]
CALIBRATED_WINDOW_WORDS = 150
DEFAULT_OVERLAP_PCT = 0
QUICK_DEFAULT_MAX_WINDOWS = 20
CALIBRATION_DEFAULT_MAX_WINDOWS = 4
PANGRAM_REALTIME_RATE_PER_100_WORDS = 0.05
PANGRAM_BULK_DISCOUNT = 0.20
PANGRAM_BULK_RATE_PER_100_WORDS = PANGRAM_REALTIME_RATE_PER_100_WORDS * (1.0 - PANGRAM_BULK_DISCOUNT)
COST_WARNING_THRESHOLD = 5.00
EXPERIMENT_WINDOW_WORDS = 500
EXPERIMENT_OVERLAP_PCT = 50
EXPERIMENT_MAX_WINDOWS_PER_FILE = 20
DEFAULT_STRUCTURE_SIMILARITY_LIMIT = 0.82
EXPERIMENT_MIN_WINDOW_RATIO = 0.90
BOOTSTRAP_PARENT_VERSION = "6E"
BOOTSTRAP_CANDIDATE_VERSION = "6F"
BOOTSTRAP_PARENT_SCREEN_AI = 0.549
BOOTSTRAP_PARENT_WHOLE_AI = 0.599
CLAUDE_DEFAULT_TARGET_WORDS = 750
CLAUDE_PANGRAM_WINDOW_WORDS = 150
CLAUDE_PANGRAM_MIN_RATIO = 0.90
CLAUDE_MAX_SCORE_WINDOWS = 20
CLAUDE_DEFAULT_MAX_TOKENS = 16000
CLAUDE_DEFAULT_EFFORT = "high"
CLAUDE_EFFORT_LEVELS = ["low", "medium", "high", "xhigh", "max"]
CLAUDE_HARNESS_STANDARD = "Standard API · v2.10 wrapper"
CLAUDE_HARNESS_PARITY = "Parity A · separate blocks, minimal instruction"
CLAUDE_HARNESS_PARITY_SYSTEM = "Parity B · separate blocks + Claude.ai system prompt"
CLAUDE_HARNESS_MODES = [
    CLAUDE_HARNESS_STANDARD,
    CLAUDE_HARNESS_PARITY,
    CLAUDE_HARNESS_PARITY_SYSTEM,
]
CLAUDE_AI_SYSTEM_PROMPT_DOCS_URL = "https://platform.claude.com/docs/en/release-notes/system-prompts"



@dataclass
class SourceDoc:
    name: str
    expected_label: str
    text: str


@dataclass
class TextWindow:
    window_id: str
    source_name: str
    expected_label: str
    target_words: int
    actual_words: int
    sentence_start: int
    sentence_end: int
    text: str


# -------------------------
# Text handling
# -------------------------

WORD_RE = re.compile(r"\S+")


def count_words(text: str) -> int:
    return len(WORD_RE.findall(text or ""))


def clean_text(text: str) -> str:
    text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    # Keep paragraph boundaries, but normalize internal whitespace.
    paragraphs = []
    for p in re.split(r"\n\s*\n", text):
        p = re.sub(r"[ \t\f\v]+", " ", p).strip()
        if p:
            paragraphs.append(p)
    return "\n\n".join(paragraphs)


def extract_text_from_upload(uploaded_file) -> str:
    suffix = Path(uploaded_file.name).suffix.lower()
    raw = uploaded_file.getvalue()

    if suffix == ".docx":
        doc = Document(io.BytesIO(raw))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return clean_text("\n\n".join(parts))

    if suffix in {".txt", ".md"}:
        for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
            try:
                return clean_text(raw.decode(encoding))
            except UnicodeDecodeError:
                continue
        return clean_text(raw.decode("utf-8", errors="replace"))

    raise ValueError(f"Unsupported file type: {suffix}. Use DOCX, TXT, or MD.")


def split_sentences(text: str) -> list[str]:
    """A lightweight fiction-friendly sentence splitter.

    Pangram 4 is designed for complete-sentence prose. We therefore build test
    windows on sentence boundaries rather than cutting at an exact word index.
    This is deliberately dependency-light; it is not intended as a linguistic
    parser.
    """
    text = clean_text(text)
    if not text:
        return []

    sentences: list[str] = []
    # Work paragraph by paragraph so paragraph breaks remain natural stopping points.
    for paragraph in re.split(r"\n\s*\n", text):
        paragraph = paragraph.strip()
        if not paragraph:
            continue

        # Capture through sentence-ending punctuation plus closing quote/bracket.
        # If a paragraph has no terminal punctuation, keep it as one unit.
        matches = re.findall(
            r".+?(?:[.!?]+(?:[\"'”’)]*)?(?=\s+|$)|$)",
            paragraph,
            flags=re.S,
        )
        for item in matches:
            item = item.strip()
            if item:
                sentences.append(item)

    return sentences


def _advance_start(sent_word_counts: list[int], start: int, stride_words: int) -> int:
    if start >= len(sent_word_counts) - 1:
        return len(sent_word_counts)
    total = 0
    i = start
    while i < len(sent_word_counts) and total < stride_words:
        total += sent_word_counts[i]
        i += 1
    return max(start + 1, i)


def build_sentence_windows(
    doc: SourceDoc,
    target_words: int,
    overlap_fraction: float,
) -> list[TextWindow]:
    sentences = split_sentences(doc.text)
    if not sentences:
        return []

    counts = [count_words(s) for s in sentences]
    stride_words = max(1, int(round(target_words * (1.0 - overlap_fraction))))
    min_acceptable = MIN_PANGRAM_WORDS

    windows: list[TextWindow] = []
    start = 0
    ordinal = 0

    while start < len(sentences):
        total = 0
        end = start
        while end < len(sentences) and total < target_words:
            total += counts[end]
            end += 1

        if total < min_acceptable:
            break

        ordinal += 1
        text = " ".join(sentences[start:end]).strip()
        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", Path(doc.name).stem)[:45]
        wid = f"{safe_name}__{target_words}w__{ordinal:04d}"
        windows.append(
            TextWindow(
                window_id=wid,
                source_name=doc.name,
                expected_label=doc.expected_label,
                target_words=target_words,
                actual_words=count_words(text),
                sentence_start=start + 1,
                sentence_end=end,
                text=text,
            )
        )

        start = _advance_start(counts, start, stride_words)

    return windows


def evenly_cap(items: list[TextWindow], cap: int) -> list[TextWindow]:
    if cap <= 0 or len(items) <= cap:
        return items
    if cap == 1:
        return [items[len(items) // 2]]
    indexes = [round(i * (len(items) - 1) / (cap - 1)) for i in range(cap)]
    seen = set()
    selected = []
    for idx in indexes:
        if idx not in seen:
            selected.append(items[idx])
            seen.add(idx)
    return selected


def make_calibration_windows(
    docs: list[SourceDoc],
    sample_sizes: list[int],
    overlap_fraction: float,
    cap_per_source_size: int,
) -> list[TextWindow]:
    out: list[TextWindow] = []
    for doc in docs:
        for size in sorted(sample_sizes):
            candidates = build_sentence_windows(doc, size, overlap_fraction)
            out.extend(evenly_cap(candidates, cap_per_source_size))
    return out


# -------------------------
# Cost estimation
# -------------------------


def estimate_bulk_cost(windows: list[TextWindow]) -> dict[str, float | int]:
    """Estimate Pangram bulk cost using started 100-word blocks per request.

    Assumption: $0.05 per started 100-word block less a 20% bulk discount,
    for an effective $0.04 per started 100-word block.
    """
    requests = len(windows)
    actual_words = sum(w.actual_words for w in windows)
    billing_blocks = sum(max(1, (w.actual_words + 99) // 100) for w in windows)
    billable_words = billing_blocks * 100
    estimated_cost = billing_blocks * PANGRAM_BULK_RATE_PER_100_WORDS
    return {
        "requests": requests,
        "actual_words": actual_words,
        "billing_blocks": billing_blocks,
        "billable_words": billable_words,
        "estimated_cost": estimated_cost,
    }


def show_cost_estimate(windows: list[TextWindow], *, key: str) -> None:
    if not windows:
        return
    est = estimate_bulk_cost(windows)
    st.info(
        f"Estimated Pangram bulk cost: **${est['estimated_cost']:.2f}** · "
        f"{est['requests']:,} requests · {est['actual_words']:,} actual words · "
        f"{est['billable_words']:,} billable words."
    )
    st.caption(
        "Estimate assumes $0.05 per started 100-word block with a 20% bulk discount "
        "($0.04 per block). Pangram bills each request separately, so sentence-aligned windows can round up."
    )
    if est["estimated_cost"] >= COST_WARNING_THRESHOLD:
        st.warning(
            f"Cost guardrail: this run is estimated at ${est['estimated_cost']:.2f}, "
            f"which is at or above the ${COST_WARNING_THRESHOLD:.2f} warning threshold."
        )


def estimate_realtime_cost(windows: list[TextWindow]) -> dict[str, float | int]:
    """Estimate Pangram realtime cost using started 100-word blocks per request."""
    requests = len(windows)
    actual_words = sum(w.actual_words for w in windows)
    billing_blocks = sum(max(1, (w.actual_words + 99) // 100) for w in windows)
    billable_words = billing_blocks * 100
    estimated_cost = billing_blocks * PANGRAM_REALTIME_RATE_PER_100_WORDS
    return {
        "requests": requests,
        "actual_words": actual_words,
        "billing_blocks": billing_blocks,
        "billable_words": billable_words,
        "estimated_cost": estimated_cost,
    }


def show_realtime_cost_estimate(windows: list[TextWindow], *, key: str) -> None:
    if not windows:
        return
    est = estimate_realtime_cost(windows)
    st.info(
        f"Estimated Pangram realtime cost: **${est['estimated_cost']:.2f}** · "
        f"{est['requests']:,} requests · {est['actual_words']:,} actual words · "
        f"{est['billable_words']:,} billable words."
    )
    st.caption(
        "Legacy A/B uses realtime requests intentionally so a two-text comparison does not sit in the lower-priority bulk queue. "
        "Estimate assumes $0.05 per started 100-word block."
    )
    if est["estimated_cost"] >= COST_WARNING_THRESHOLD:
        st.warning(
            f"Cost guardrail: this run is estimated at ${est['estimated_cost']:.2f}, "
            f"which is at or above the ${COST_WARNING_THRESHOLD:.2f} warning threshold."
        )


# -------------------------
# Anthropic API
# -------------------------


def get_anthropic_api_key() -> str:
    env_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if env_key:
        return env_key
    try:
        return str(st.secrets.get("ANTHROPIC_API_KEY", "")).strip()
    except Exception:
        return ""


def connect_anthropic(api_key: str) -> tuple[Any, list[str]]:
    if Anthropic is None:
        raise RuntimeError(
            "The anthropic package is not installed. Add anthropic to requirements.txt, then reboot the Streamlit app."
        )
    client = Anthropic(api_key=api_key, timeout=300.0, max_retries=1)
    page = client.models.list(limit=100)
    model_ids = [str(m.id) for m in getattr(page, "data", []) if getattr(m, "id", None)]
    return client, model_ids


def preferred_anthropic_model(models: list[str]) -> str | None:
    for preferred in (
        "claude-opus-5",
        "claude-sonnet-5",
        "claude-opus-4-8",
        "claude-sonnet-4-6",
    ):
        if preferred in models:
            return preferred
    return models[0] if models else None


def get_connected_anthropic() -> tuple[Any | None, str | None]:
    return st.session_state.get("anthropic_client"), st.session_state.get("anthropic_model")


def detect_word_count_target(text: str) -> int | None:
    patterns = [
        r"word\s+count\s+target\s*:\s*([0-9][0-9,]*)\s*words?",
        r"target\s+word\s+count\s*:\s*([0-9][0-9,]*)\s*words?",
        r"approximately\s+([0-9][0-9,]*)\s+words?",
    ]
    for pattern in patterns:
        match = re.search(pattern, text or "", flags=re.I)
        if match:
            try:
                return int(match.group(1).replace(",", ""))
            except Exception:
                pass
    return None


def build_claude_test_message(prompt_text: str, packet_text: str, target_words: int) -> str:
    """Create one fixed, auditable API payload. No prior runs or detector feedback enter it."""
    return (
        "<drafting_instructions>\n"
        + prompt_text.strip()
        + "\n</drafting_instructions>\n\n"
        + "<chapter_packet>\n"
        + packet_text.strip()
        + "\n</chapter_packet>\n\n"
        + "<rapid_test_harness>\n"
        + f"Draft the chapter now. For this rapid test, aim for approximately {int(target_words)} words. "
          "For this API test, this harness overrides any output-format instruction inside the drafting instructions. "
          "Return only the finished fiction prose as plain text. Do not emit XML, Markdown wrappers, artifact markup, "
          "tool calls, document-creation commands, a title, a preface, an explanation, analysis, a checklist, notes, "
          "or a self-critique. Do not describe your compliance. The application will create the DOCX download.\n"
        + "</rapid_test_harness>"
    )


def build_claude_request(
    prompt_text: str,
    packet_text: str,
    target_words: int,
    harness_mode: str,
    claude_ai_system_prompt: str = "",
) -> dict[str, Any]:
    """Build an auditable request while keeping the two source documents distinct in parity modes.

    Standard mode preserves the v2.10 wrapper exactly.  Parity modes remove that wrapper,
    preserve the prompt and packet as separate text content blocks, and add only the minimum
    instruction required for an API response instead of a web-created DOCX artifact.
    """
    if harness_mode == CLAUDE_HARNESS_STANDARD:
        payload = build_claude_test_message(prompt_text, packet_text, int(target_words))
        return {
            "messages": [{"role": "user", "content": payload}],
            "system": None,
            "description": "Standard v2.10 wrapper",
            "preview": payload,
        }

    content_blocks = [
        {"type": "text", "text": prompt_text.strip()},
        {"type": "text", "text": packet_text.strip()},
        {
            "type": "text",
            "text": (
                "Write the chapter now. Return only the finished chapter prose in this response. "
                "The application will save it as a DOCX."
            ),
        },
    ]
    system_prompt = None
    description = "Parity A: separate source blocks; minimal final instruction"
    if harness_mode == CLAUDE_HARNESS_PARITY_SYSTEM:
        system_prompt = (claude_ai_system_prompt or "").strip() or None
        description = "Parity B: separate source blocks + supplied Claude.ai system prompt"

    preview_parts = []
    if system_prompt:
        preview_parts.append("SYSTEM PROMPT\n" + system_prompt)
    for i, block in enumerate(content_blocks, start=1):
        preview_parts.append(f"USER CONTENT BLOCK {i}\n{block['text']}")

    return {
        "messages": [{"role": "user", "content": content_blocks}],
        "system": system_prompt,
        "description": description,
        "preview": "\n\n".join(preview_parts),
    }


def run_claude_stream(
    client: Any,
    model: str,
    messages: list[dict[str, Any]],
    *,
    max_tokens: int,
    effort: str,
    live_placeholder: Any,
    system_prompt: str | None = None,
) -> tuple[str, Any, float, str]:
    kwargs: dict[str, Any] = {
        "model": model,
        "max_tokens": int(max_tokens),
        "messages": messages,
    }
    if system_prompt:
        kwargs["system"] = system_prompt

    # Claude 5 models use adaptive thinking by default when thinking is omitted.
    # Keep thinking omitted here so parity mode stays closer to claude.ai.
    current_effort_models = {"claude-fable-5", "claude-opus-5", "claude-sonnet-5"}
    if model in current_effort_models:
        kwargs["output_config"] = {"effort": str(effort)}
        generation_mode = f"Adaptive/default · effort {effort}"
    else:
        generation_mode = "Model default · effort omitted"

    started = time.perf_counter()
    chunks: list[str] = []
    last_paint = 0.0
    with client.messages.stream(**kwargs) as stream:
        for piece in stream.text_stream:
            chunks.append(piece)
            now = time.perf_counter()
            if now - last_paint >= 0.25:
                live_placeholder.markdown("".join(chunks))
                last_paint = now
        final_message = stream.get_final_message()
    elapsed = time.perf_counter() - started

    final_text = "".join(chunks).strip()
    if not final_text:
        text_blocks = [
            getattr(block, "text", "")
            for block in getattr(final_message, "content", [])
            if getattr(block, "type", None) == "text"
        ]
        final_text = "".join(text_blocks).strip()
    live_placeholder.markdown(final_text or "*(No text returned.)*")
    return final_text, final_message, elapsed, generation_mode


def claude_stop_details(final_message: Any) -> tuple[str | None, str | None, str | None]:
    """Return stop_reason, refusal category, and human-readable explanation when present."""
    stop_reason = getattr(final_message, "stop_reason", None)
    details = getattr(final_message, "stop_details", None)
    category = getattr(details, "category", None) if details is not None else None
    explanation = getattr(details, "explanation", None) if details is not None else None
    return stop_reason, category, explanation


def prose_docx_bytes(text: str) -> bytes:
    doc = Document()
    for paragraph in re.split(r"\n\s*\n", clean_text(text)):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generation_score_summary(rows: list[dict[str, Any]]) -> dict[str, Any]:
    df = results_dataframe(rows)
    if df.empty:
        return {
            "mean_ai": None,
            "mean_involvement": None,
            "human_windows": 0,
            "mixed_windows": 0,
            "ai_windows": 0,
        }
    weights = df["actual_words"].fillna(1).astype(float)
    denom = float(weights.sum()) or 1.0
    mean_ai = float((df["fraction_ai"].fillna(0).astype(float) * weights).sum() / denom)
    mean_involvement = float((df["mean_ai_involvement"].fillna(0).astype(float) * weights).sum() / denom)
    counts = df["prediction"].value_counts().to_dict()
    return {
        "mean_ai": mean_ai,
        "mean_involvement": mean_involvement,
        "human_windows": int(counts.get("Human", 0)),
        "mixed_windows": int(counts.get("Mixed", 0)),
        "ai_windows": int(counts.get("AI", 0)),
    }


# -------------------------
# Pangram API
# -------------------------


def get_secret_api_key() -> str:
    env_key = os.getenv("PANGRAM_API_KEY", "").strip()
    if env_key:
        return env_key
    try:
        return str(st.secrets.get("PANGRAM_API_KEY", "")).strip()
    except Exception:
        return ""


def connect_pangram(api_key: str) -> tuple[Any, list[str]]:
    if Pangram is None:
        raise RuntimeError(
            "The pangram-sdk package is not installed. Make sure requirements.txt is in the GitHub repo, then reboot the Streamlit app."
        )
    client = Pangram(api_key=api_key) if api_key else Pangram()
    models = client.list_models()
    return client, list(models)


def weighted_window_metric(result: dict[str, Any], field: str) -> float | None:
    windows = result.get("windows") or []
    vals = []
    weights = []
    for w in windows:
        value = w.get(field)
        if value is None:
            continue
        weight = w.get("word_count") or count_words(w.get("text", "")) or 1
        vals.append(float(value))
        weights.append(float(weight))
    if not vals:
        return None
    return sum(v * wt for v, wt in zip(vals, weights)) / sum(weights)


def summarize_result(result: dict[str, Any]) -> dict[str, Any]:
    windows = result.get("windows") or []
    confidences = [str(w.get("confidence", "")) for w in windows if w.get("confidence")]
    humanizer_scores = [
        float(w["humanizer_score"])
        for w in windows
        if w.get("humanizer_score") is not None
    ]
    return {
        "prediction": result.get("prediction_short"),
        "headline": result.get("headline"),
        "fraction_ai": result.get("fraction_ai"),
        "fraction_ai_assisted": result.get("fraction_ai_assisted"),
        "fraction_human": result.get("fraction_human"),
        "num_ai_segments": result.get("num_ai_segments"),
        "num_ai_assisted_segments": result.get("num_ai_assisted_segments"),
        "num_human_segments": result.get("num_human_segments"),
        "mean_ai_involvement": weighted_window_metric(result, "ai_assistance_score"),
        "max_humanizer_score": max(humanizer_scores) if humanizer_scores else None,
        "any_humanized": any(bool(w.get("is_humanized")) for w in windows),
        "window_confidence": ", ".join(sorted(set(confidences))),
        "version": result.get("version"),
        "dashboard_link": result.get("dashboard_link"),
    }


def run_realtime_scan(
    client: Any,
    model: str,
    windows: list[TextWindow],
    timeout: float = 300,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Run small paired tests as individual realtime Pangram requests.

    This avoids sending tiny A/B experiments into the asynchronous bulk queue,
    where a two-item job can remain pending for a long time.
    """
    successes: list[dict[str, Any]] = []
    failures: list[dict[str, Any]] = []
    total = len(windows)
    progress = st.progress(0, text="Starting Pangram realtime scan…")

    for i, w in enumerate(windows, start=1):
        progress.progress(
            (i - 1) / max(1, total),
            text=f"Pangram realtime request {i}/{total}…",
        )
        try:
            result = client.predict(
                w.text,
                model=model,
                timeout=timeout,
                poll_interval=0.5,
            )
            row = asdict(w)
            row.update(summarize_result(result))
            row["bulk_id"] = None
            row["raw_result"] = result
            successes.append(row)
        except Exception as exc:
            failures.append(
                {
                    "window_id": w.window_id,
                    "error": str(exc),
                    "bulk_id": None,
                }
            )

    progress.progress(1.0, text="Pangram realtime analysis complete.")
    time.sleep(0.15)
    progress.empty()
    return successes, failures


def run_bulk_scan(
    client: Any,
    model: str,
    windows: list[TextWindow],
    batch_size: int = 200,
    timeout: float = 3600,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    """Submit one or more Pangram bulk jobs and return successful and failed rows."""
    successes: list[dict[str, Any]] = []
    failures: list[dict[str, Any]] = []

    meta_by_id = {w.window_id: w for w in windows}
    progress = st.progress(0, text="Submitting Pangram bulk job…")
    total_batches = max(1, (len(windows) + batch_size - 1) // batch_size)

    for batch_no, offset in enumerate(range(0, len(windows), batch_size), start=1):
        chunk = windows[offset : offset + batch_size]
        items = [{"id": w.window_id, "text": w.text} for w in chunk]

        bulk = client.submit_bulk(items=items, model=model)
        bulk_id = bulk["bulk_id"]
        progress.progress(
            min(0.85, (batch_no - 1) / total_batches + 0.05),
            text=f"Pangram batch {batch_no}/{total_batches}: waiting for results…",
        )
        status = client.wait_for_bulk(bulk_id, timeout=timeout, poll_interval=0.5)
        results = client.get_bulk_results(bulk_id)

        for item in results.get("items", []):
            item_id = item.get("id")
            result = item.get("result")
            if result is None:
                failures.append(
                    {
                        "window_id": item_id,
                        "error": item.get("error") or f"No result; stage={item.get('stage')}",
                        "bulk_id": bulk_id,
                    }
                )
                continue

            meta = meta_by_id.get(item_id)
            if meta is None:
                failures.append(
                    {"window_id": item_id, "error": "Unknown result ID", "bulk_id": bulk_id}
                )
                continue

            row = asdict(meta)
            row.update(summarize_result(result))
            row["bulk_id"] = bulk_id
            row["raw_result"] = result
            successes.append(row)

        for failed in results.get("failed_items", []):
            failures.append(
                {
                    "window_id": failed.get("id"),
                    "error": failed.get("error") or "Pangram bulk item failed",
                    "bulk_id": bulk_id,
                }
            )

        progress.progress(
            min(0.98, batch_no / total_batches),
            text=f"Pangram batch {batch_no}/{total_batches} complete.",
        )

    progress.progress(1.0, text="Pangram analysis complete.")
    time.sleep(0.15)
    progress.empty()
    return successes, failures


# -------------------------
# Persistence
# -------------------------


def _ensure_column(con: sqlite3.Connection, table: str, column: str, declaration: str) -> None:
    existing = {row[1] for row in con.execute(f"PRAGMA table_info({table})")}
    if column not in existing:
        con.execute(f"ALTER TABLE {table} ADD COLUMN {column} {declaration}")


def init_db() -> None:
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS scan_results (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                experiment_id TEXT NOT NULL,
                experiment_name TEXT,
                run_at TEXT NOT NULL,
                mode TEXT NOT NULL,
                model TEXT,
                source_name TEXT,
                expected_label TEXT,
                window_id TEXT,
                target_words INTEGER,
                actual_words INTEGER,
                sentence_start INTEGER,
                sentence_end INTEGER,
                prediction TEXT,
                headline TEXT,
                fraction_ai REAL,
                fraction_ai_assisted REAL,
                fraction_human REAL,
                mean_ai_involvement REAL,
                max_humanizer_score REAL,
                any_humanized INTEGER,
                window_confidence TEXT,
                version TEXT,
                dashboard_link TEXT,
                text TEXT,
                raw_json TEXT
            )
            """
        )
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS experiment_runs (
                experiment_id TEXT PRIMARY KEY,
                run_at TEXT NOT NULL,
                parent_version TEXT,
                candidate_version TEXT,
                change_note TEXT,
                test_set_note TEXT,
                parent_prompt TEXT,
                candidate_prompt TEXT,
                model TEXT,
                target_words INTEGER,
                overlap_pct INTEGER,
                max_windows_per_file INTEGER,
                parent_files INTEGER,
                candidate_files INTEGER,
                parent_mean_ai REAL,
                candidate_mean_ai REAL,
                delta_ai REAL,
                candidate_worst_ai REAL,
                candidate_max_structure_similarity REAL,
                structure_similarity_limit REAL,
                verdict TEXT
            )
            """
        )
        for column, declaration in [
            ("parent_whole_ai", "REAL"),
            ("candidate_whole_ai", "REAL"),
            ("whole_delta_ai", "REAL"),
            ("candidate_main_windows", "INTEGER"),
            ("candidate_tail_windows", "INTEGER"),
            ("min_window_ratio", "REAL"),
        ]:
            _ensure_column(con, "experiment_runs", column, declaration)

        con.execute(
            """
            CREATE TABLE IF NOT EXISTS baselines (
                version TEXT PRIMARY KEY,
                saved_at TEXT NOT NULL,
                screen_ai REAL,
                screen_ai_involvement REAL,
                whole_ai REAL,
                source_note TEXT,
                prompt_text TEXT
            )
            """
        )
        con.execute(
            """
            CREATE TABLE IF NOT EXISTS generation_runs (
                run_id TEXT PRIMARY KEY,
                run_at TEXT NOT NULL,
                label TEXT,
                anthropic_model TEXT,
                thinking_mode TEXT,
                target_words INTEGER,
                output_words INTEGER,
                max_tokens INTEGER,
                input_tokens INTEGER,
                output_tokens INTEGER,
                stop_reason TEXT,
                generation_seconds REAL,
                pangram_model TEXT,
                pangram_seconds REAL,
                pangram_windows INTEGER,
                mean_ai REAL,
                mean_ai_involvement REAL,
                prompt_name TEXT,
                packet_name TEXT,
                prompt_text TEXT,
                packet_text TEXT,
                output_text TEXT
            )
            """
        )
        con.commit()


def save_results(
    rows: list[dict[str, Any]],
    *,
    experiment_name: str,
    mode: str,
    model: str,
    experiment_id: str | None = None,
) -> str:
    init_db()
    experiment_id = experiment_id or str(uuid.uuid4())
    run_at = datetime.now(timezone.utc).isoformat()

    with sqlite3.connect(DB_PATH) as con:
        for r in rows:
            con.execute(
                """
                INSERT INTO scan_results (
                    experiment_id, experiment_name, run_at, mode, model,
                    source_name, expected_label, window_id, target_words, actual_words,
                    sentence_start, sentence_end, prediction, headline,
                    fraction_ai, fraction_ai_assisted, fraction_human,
                    mean_ai_involvement, max_humanizer_score, any_humanized,
                    window_confidence, version, dashboard_link, text, raw_json
                ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    experiment_id,
                    experiment_name,
                    run_at,
                    mode,
                    model,
                    r.get("source_name"),
                    r.get("expected_label"),
                    r.get("window_id"),
                    r.get("target_words"),
                    r.get("actual_words"),
                    r.get("sentence_start"),
                    r.get("sentence_end"),
                    r.get("prediction"),
                    r.get("headline"),
                    r.get("fraction_ai"),
                    r.get("fraction_ai_assisted"),
                    r.get("fraction_human"),
                    r.get("mean_ai_involvement"),
                    r.get("max_humanizer_score"),
                    1 if r.get("any_humanized") else 0,
                    r.get("window_confidence"),
                    r.get("version"),
                    r.get("dashboard_link"),
                    r.get("text"),
                    json.dumps(r.get("raw_result") or {}, ensure_ascii=False),
                ),
            )
        con.commit()
    return experiment_id


def save_experiment_run(record: dict[str, Any]) -> None:
    init_db()
    cols = [
        "experiment_id", "run_at", "parent_version", "candidate_version",
        "change_note", "test_set_note", "parent_prompt", "candidate_prompt",
        "model", "target_words", "overlap_pct", "max_windows_per_file",
        "parent_files", "candidate_files", "parent_mean_ai", "candidate_mean_ai",
        "delta_ai", "candidate_worst_ai", "candidate_max_structure_similarity",
        "structure_similarity_limit", "verdict", "parent_whole_ai",
        "candidate_whole_ai", "whole_delta_ai", "candidate_main_windows",
        "candidate_tail_windows", "min_window_ratio"
    ]
    values = [record.get(c) for c in cols]
    placeholders = ",".join(["?"] * len(cols))
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            f"INSERT OR REPLACE INTO experiment_runs ({','.join(cols)}) VALUES ({placeholders})",
            values,
        )
        con.commit()


def save_generation_run(record: dict[str, Any]) -> None:
    init_db()
    cols = [
        "run_id", "run_at", "label", "anthropic_model", "thinking_mode",
        "target_words", "output_words", "max_tokens", "input_tokens",
        "output_tokens", "stop_reason", "generation_seconds", "pangram_model",
        "pangram_seconds", "pangram_windows", "mean_ai", "mean_ai_involvement",
        "prompt_name", "packet_name", "prompt_text", "packet_text", "output_text",
    ]
    placeholders = ",".join(["?"] * len(cols))
    values = [record.get(c) for c in cols]
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            f"INSERT OR REPLACE INTO generation_runs ({','.join(cols)}) VALUES ({placeholders})",
            values,
        )
        con.commit()


def load_generation_runs(limit: int = 100) -> pd.DataFrame:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        return pd.read_sql_query(
            "SELECT * FROM generation_runs ORDER BY run_at DESC LIMIT ?",
            con,
            params=(int(limit),),
        )


def load_history() -> pd.DataFrame:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        return pd.read_sql_query(
            "SELECT * FROM scan_results ORDER BY id DESC",
            con,
        )


def load_experiment_history() -> pd.DataFrame:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        return pd.read_sql_query(
            "SELECT * FROM experiment_runs ORDER BY run_at DESC",
            con,
        )


def save_baseline(
    version: str,
    *,
    screen_ai: float | None,
    screen_ai_involvement: float | None,
    whole_ai: float | None,
    source_note: str = "",
    prompt_text: str = "",
) -> None:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        con.execute(
            """
            INSERT OR REPLACE INTO baselines
            (version, saved_at, screen_ai, screen_ai_involvement, whole_ai, source_note, prompt_text)
            VALUES (?,?,?,?,?,?,?)
            """,
            (
                version.strip(),
                datetime.now(timezone.utc).isoformat(),
                screen_ai,
                screen_ai_involvement,
                whole_ai,
                source_note,
                prompt_text,
            ),
        )
        con.commit()


def load_baselines() -> pd.DataFrame:
    init_db()
    with sqlite3.connect(DB_PATH) as con:
        return pd.read_sql_query(
            "SELECT * FROM baselines ORDER BY saved_at DESC",
            con,
        )


# -------------------------
# Analysis / display
# -------------------------


def results_dataframe(rows: list[dict[str, Any]]) -> pd.DataFrame:
    if not rows:
        return pd.DataFrame()
    display_rows = []
    for r in rows:
        d = {k: v for k, v in r.items() if k != "raw_result"}
        display_rows.append(d)
    return pd.DataFrame(display_rows)


def calibration_summary(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return df
    usable = df[df["expected_label"].isin(["Human", "AI"])].copy()
    if usable.empty:
        return pd.DataFrame()

    records = []
    for (label, size), group in usable.groupby(["expected_label", "target_words"]):
        n = len(group)
        records.append(
            {
                "Expected": label,
                "Target words": int(size),
                "N": n,
                "Human %": 100 * (group["prediction"] == "Human").mean(),
                "Mixed %": 100 * (group["prediction"] == "Mixed").mean(),
                "AI %": 100 * (group["prediction"] == "AI").mean(),
                "Mean human fraction": group["fraction_human"].mean(),
                "Mean AI fraction": group["fraction_ai"].mean(),
                "Mean AI involvement": group["mean_ai_involvement"].mean(),
            }
        )
    return pd.DataFrame(records).sort_values(["Target words", "Expected"])


def separation_summary(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    records = []
    for size, group in df.groupby("target_words"):
        human = group[group["expected_label"] == "Human"]
        ai = group[group["expected_label"] == "AI"]
        if human.empty or ai.empty:
            continue
        human_correct = (human["prediction"] == "Human").mean()
        ai_correct = (ai["prediction"] == "AI").mean()
        records.append(
            {
                "Target words": int(size),
                "Human correctly Human %": 100 * human_correct,
                "AI correctly AI %": 100 * ai_correct,
                "Balanced decisive accuracy %": 100 * (human_correct + ai_correct) / 2,
                "Human N": len(human),
                "AI N": len(ai),
            }
        )
    return pd.DataFrame(records).sort_values("Target words")


def recommend_size(sep: pd.DataFrame, human_threshold: float, ai_threshold: float) -> str:
    if sep.empty:
        return "Not enough labeled Human and AI data to recommend a window size."
    candidates = sep[
        (sep["Human correctly Human %"] >= human_threshold * 100)
        & (sep["AI correctly AI %"] >= ai_threshold * 100)
    ]
    if candidates.empty:
        return (
            "No tested size met both thresholds. That is useful: either test larger windows, "
            "add more control texts, or accept a lower screening threshold for early experiments."
        )
    size = int(candidates.iloc[0]["Target words"])
    return (
        f"Smallest tested size meeting both thresholds: about **{size} words** "
        "(sentence-aligned, so individual windows may be somewhat longer)."
    )


def show_result_table(df: pd.DataFrame, key: str) -> None:
    if df.empty:
        st.info("No results yet.")
        return
    preferred = [
        "source_name",
        "expected_label",
        "target_words",
        "actual_words",
        "prediction",
        "fraction_human",
        "fraction_ai_assisted",
        "fraction_ai",
        "mean_ai_involvement",
        "max_humanizer_score",
        "window_confidence",
        "sentence_start",
        "sentence_end",
        "text",
    ]
    cols = [c for c in preferred if c in df.columns]
    st.dataframe(df[cols], use_container_width=True, hide_index=True, key=key)


def source_docs_from_uploads(files: Iterable[Any], expected_label: str) -> tuple[list[SourceDoc], list[str]]:
    docs = []
    errors = []
    for f in files or []:
        try:
            text = extract_text_from_upload(f)
            if count_words(text) < MIN_PANGRAM_WORDS:
                errors.append(f"{f.name}: fewer than {MIN_PANGRAM_WORDS} words after extraction")
                continue
            docs.append(SourceDoc(name=f.name, expected_label=expected_label, text=text))
        except Exception as exc:
            errors.append(f"{f.name}: {exc}")
    return docs, errors


def get_connected_client() -> tuple[Any | None, str | None]:
    return st.session_state.get("pangram_client"), st.session_state.get("pangram_model")


# -------------------------
# Experiment Lab analysis
# -------------------------


def weighted_mean(values: pd.Series, weights: pd.Series) -> float | None:
    usable = pd.DataFrame({"v": values, "w": weights}).dropna()
    if usable.empty or usable["w"].sum() <= 0:
        return None
    return float((usable["v"] * usable["w"]).sum() / usable["w"].sum())


def experiment_file_summary(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    records = []
    for (side, source_name), group in df.groupby(["expected_label", "source_name"], sort=False):
        counts = group["prediction"].value_counts().to_dict()
        records.append(
            {
                "Version": side,
                "File": source_name,
                "Windows": len(group),
                "Submitted words": int(group["actual_words"].sum()),
                "Weighted AI fraction": weighted_mean(group["fraction_ai"], group["actual_words"]),
                "Weighted AI involvement": weighted_mean(group["mean_ai_involvement"], group["actual_words"]),
                "Human %": 100 * counts.get("Human", 0) / len(group),
                "Mixed %": 100 * counts.get("Mixed", 0) / len(group),
                "AI %": 100 * counts.get("AI", 0) / len(group),
            }
        )
    return pd.DataFrame(records)


def side_summary(file_summary: pd.DataFrame) -> pd.DataFrame:
    if file_summary.empty:
        return pd.DataFrame()
    records = []
    for version, group in file_summary.groupby("Version", sort=False):
        vals = group["Weighted AI fraction"].dropna()
        inv = group["Weighted AI involvement"].dropna()
        records.append(
            {
                "Version": version,
                "Files": len(group),
                "Mean AI fraction": float(vals.mean()) if not vals.empty else None,
                "Worst-file AI fraction": float(vals.max()) if not vals.empty else None,
                "Best-file AI fraction": float(vals.min()) if not vals.empty else None,
                "AI-fraction stdev": float(vals.std(ddof=0)) if len(vals) > 1 else 0.0,
                "Mean AI involvement": float(inv.mean()) if not inv.empty else None,
                "Human windows %": float(group["Human %"].mean()),
                "Mixed windows %": float(group["Mixed %"].mean()),
                "AI windows %": float(group["AI %"].mean()),
            }
        )
    return pd.DataFrame(records)


def _bucket_word_count(n: int) -> str:
    if n <= 25:
        return "A"
    if n <= 50:
        return "B"
    if n <= 90:
        return "C"
    if n <= 140:
        return "D"
    return "E"


def _bucket_sentence_count(n: int) -> str:
    return str(n) if n <= 4 else "5+"


def _bucket_sentence_words(n: int) -> str:
    if n <= 7:
        return "XS"
    if n <= 14:
        return "S"
    if n <= 24:
        return "M"
    if n <= 40:
        return "L"
    return "XL"


def structure_signature(text: str) -> dict[str, Any]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", clean_text(text)) if p.strip()]
    para_tokens: list[str] = []
    sentence_tokens: list[str] = []
    for p in paragraphs:
        sents = split_sentences(p)
        dialogue = p.lstrip().startswith(('"', '“', "'", '‘'))
        para_tokens.append(
            f"{_bucket_word_count(count_words(p))}|{_bucket_sentence_count(max(1, len(sents)))}|{'D' if dialogue else 'N'}"
        )
        for sent in sents:
            sentence_tokens.append(_bucket_sentence_words(count_words(sent)))
    return {
        "paragraphs": len(paragraphs),
        "sentences": len(sentence_tokens),
        "para_tokens": para_tokens,
        "sentence_tokens": sentence_tokens,
    }


def structure_similarity(text_a: str, text_b: str) -> dict[str, float]:
    a = structure_signature(text_a)
    b = structure_signature(text_b)
    para = difflib.SequenceMatcher(None, a["para_tokens"], b["para_tokens"], autojunk=False).ratio()
    sent = difflib.SequenceMatcher(None, a["sentence_tokens"], b["sentence_tokens"], autojunk=False).ratio()
    combined = 0.70 * para + 0.30 * sent
    return {
        "paragraph_similarity": float(para),
        "sentence_shape_similarity": float(sent),
        "combined_similarity": float(combined),
    }


def candidate_diversity_table(docs: list[SourceDoc]) -> pd.DataFrame:
    if len(docs) < 2:
        return pd.DataFrame()
    rows = []
    for i in range(len(docs)):
        for j in range(i + 1, len(docs)):
            sim = structure_similarity(docs[i].text, docs[j].text)
            rows.append(
                {
                    "File A": docs[i].name,
                    "File B": docs[j].name,
                    "Paragraph-shape similarity": sim["paragraph_similarity"],
                    "Sentence-shape similarity": sim["sentence_shape_similarity"],
                    "Combined structural similarity": sim["combined_similarity"],
                }
            )
    return pd.DataFrame(rows).sort_values("Combined structural similarity", ascending=False)


def build_experiment_windows(
    docs: list[SourceDoc],
    target_words: int,
    overlap_pct: int,
    cap_per_file: int,
) -> list[TextWindow]:
    out: list[TextWindow] = []
    for doc in docs:
        out.extend(
            evenly_cap(
                build_sentence_windows(doc, target_words, overlap_pct / 100.0),
                cap_per_file,
            )
        )
    return out


def split_experiment_windows(
    docs: list[SourceDoc],
    target_words: int,
    overlap_pct: int,
    cap_per_file: int,
    min_ratio: float = EXPERIMENT_MIN_WINDOW_RATIO,
) -> tuple[list[TextWindow], list[TextWindow]]:
    """Return score-eligible windows and undersized tail fragments.

    v2.1 deliberately excludes windows below min_ratio * target_words from the
    main score and from the default Pangram submission. This prevents short
    end-of-document fragments from disproportionately improving or worsening
    the experiment summary.
    """
    main: list[TextWindow] = []
    tails: list[TextWindow] = []
    threshold = max(MIN_PANGRAM_WORDS, int(round(target_words * min_ratio)))
    for doc in docs:
        candidates = build_sentence_windows(doc, target_words, overlap_pct / 100.0)
        doc_main = [w for w in candidates if w.actual_words >= threshold]
        doc_tails = [w for w in candidates if w.actual_words < threshold]
        main.extend(evenly_cap(doc_main, cap_per_file))
        tails.extend(doc_tails)
    return main, tails


def build_whole_document_windows(docs: list[SourceDoc]) -> list[TextWindow]:
    out: list[TextWindow] = []
    for ordinal, doc in enumerate(docs, start=1):
        wc = count_words(doc.text)
        if wc < MIN_PANGRAM_WORDS:
            continue
        safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", Path(doc.name).stem)[:45]
        out.append(
            TextWindow(
                window_id=f"{safe_name}__whole__{ordinal:04d}",
                source_name=doc.name,
                expected_label=doc.expected_label,
                target_words=wc,
                actual_words=wc,
                sentence_start=1,
                sentence_end=len(split_sentences(doc.text)),
                text=doc.text,
            )
        )
    return out


def whole_document_summary(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    records = []
    for (version, source_name), group in df.groupby(["expected_label", "source_name"], sort=False):
        r = group.iloc[0]
        records.append(
            {
                "Version": version,
                "File": source_name,
                "Words": int(r["actual_words"]),
                "Prediction": r.get("prediction"),
                "AI fraction": float(r["fraction_ai"]) if pd.notna(r.get("fraction_ai")) else None,
                "AI involvement": float(r["mean_ai_involvement"]) if pd.notna(r.get("mean_ai_involvement")) else None,
                "Human fraction": float(r["fraction_human"]) if pd.notna(r.get("fraction_human")) else None,
            }
        )
    return pd.DataFrame(records)


def tail_window_table(tails: list[TextWindow], target_words: int) -> pd.DataFrame:
    if not tails:
        return pd.DataFrame()
    return pd.DataFrame(
        [
            {
                "File": w.source_name,
                "Words": w.actual_words,
                "Target": target_words,
                "Target %": w.actual_words / target_words if target_words else None,
                "Sentence start": w.sentence_start,
                "Sentence end": w.sentence_end,
            }
            for w in tails
        ]
    )


def mean_whole_ai(whole_df: pd.DataFrame) -> float | None:
    if whole_df.empty or "AI fraction" not in whole_df:
        return None
    vals = pd.to_numeric(whole_df["AI fraction"], errors="coerce").dropna()
    return float(vals.mean()) if not vals.empty else None


def prompt_diff(parent_prompt: str, candidate_prompt: str) -> str:
    if not parent_prompt.strip() and not candidate_prompt.strip():
        return ""
    diff = difflib.unified_diff(
        parent_prompt.splitlines(),
        candidate_prompt.splitlines(),
        fromfile="parent prompt",
        tofile="candidate prompt",
        lineterm="",
    )
    return "\n".join(diff)


def experiment_verdict(delta_ai: float | None) -> str:
    if delta_ai is None:
        return "NO COMPARISON"
    if delta_ai <= -0.10:
        return "STRONG IMPROVEMENT"
    if delta_ai <= -0.05:
        return "PROMISING"
    if delta_ai < 0.05:
        return "INCONCLUSIVE"
    return "WORSE"


def compact_handoff(
    parent_version: str,
    candidate_version: str,
    change_note: str,
    file_summary: pd.DataFrame,
    side_summary_df: pd.DataFrame,
    diversity_df: pd.DataFrame,
    structure_limit: float,
) -> str:
    lines = [
        f"Pangram Experiment Lab: {parent_version} → {candidate_version}",
        f"Change tested: {change_note or '(not entered)'}",
        "",
        "Version summary:",
    ]
    if not side_summary_df.empty:
        for _, r in side_summary_df.iterrows():
            lines.append(
                f"- {r['Version']}: {int(r['Files'])} file(s), mean AI fraction {100*r['Mean AI fraction']:.1f}%, "
                f"worst file {100*r['Worst-file AI fraction']:.1f}%"
            )
    if not file_summary.empty:
        lines.append("")
        lines.append("Files:")
        for _, r in file_summary.iterrows():
            lines.append(
                f"- {r['Version']} / {r['File']}: {r['Windows']} windows, weighted AI {100*r['Weighted AI fraction']:.1f}%"
            )
    if not diversity_df.empty:
        max_sim = float(diversity_df["Combined structural similarity"].max())
        lines += [
            "",
            f"Candidate max pairwise structural similarity: {100*max_sim:.1f}% "
            f"(experimental warning line {100*structure_limit:.0f}%).",
        ]
    else:
        lines += ["", "Candidate structural diversity: not testable from fewer than 2 candidate files."]
    return "\n".join(lines)


# -------------------------
# Streamlit app
# -------------------------

st.set_page_config(page_title=APP_TITLE, layout="wide")
init_db()

st.title(f"{APP_TITLE} {APP_VERSION}")
st.caption(
    "Generate a clean, stateless Claude API sample and score it immediately with Pangram, while keeping "
    "the existing calibration, prompt-experiment, and structural-diversity tools."
)

with st.sidebar:
    st.header("Pangram connection")
    api_key = get_secret_api_key()

    if Pangram is None:
        st.error(
            "pangram-sdk is not installed. Make sure requirements.txt is in the GitHub repo, "
            "then reboot the Streamlit app."
        )
    elif not api_key:
        st.error("PANGRAM_API_KEY is not set in Streamlit Secrets.")
        st.caption("In Streamlit Cloud: Manage app → Settings → Secrets, then add:")
        st.code('PANGRAM_API_KEY = "paste-your-key-here"', language="toml")
        st.caption("Save the secret and let Streamlit rerun the app. Do not put the key in GitHub.")
    else:
        # Auto-connect once per Streamlit session. No command line and no key-pasting in the app.
        if st.session_state.get("pangram_client") is None:
            try:
                with st.spinner("Connecting to Pangram…"):
                    client, models = connect_pangram(api_key)
                st.session_state["pangram_client"] = client
                st.session_state["pangram_models"] = models
                if "pangram-4" in models:
                    st.session_state["pangram_model"] = "pangram-4"
                elif models:
                    st.session_state["pangram_model"] = models[0]
            except Exception as exc:
                st.session_state.pop("pangram_client", None)
                st.session_state.pop("pangram_models", None)
                st.session_state.pop("pangram_model", None)
                st.error(f"Pangram connection failed: {exc}")

        if st.session_state.get("pangram_client") is not None:
            st.success("API key loaded from Streamlit Secrets.")

            if st.button("Refresh Pangram models", use_container_width=True):
                try:
                    with st.spinner("Refreshing Pangram models…"):
                        client, models = connect_pangram(api_key)
                    st.session_state["pangram_client"] = client
                    st.session_state["pangram_models"] = models
                    current = st.session_state.get("pangram_model")
                    if current not in models:
                        st.session_state["pangram_model"] = (
                            "pangram-4" if "pangram-4" in models else (models[0] if models else None)
                        )
                    st.rerun()
                except Exception as exc:
                    st.error(f"Could not refresh models: {exc}")

    models = st.session_state.get("pangram_models", [])
    if models:
        current_model = st.session_state.get("pangram_model")
        idx = models.index(current_model) if current_model in models else 0
        selected = st.selectbox("Model", models, index=idx)
        st.session_state["pangram_model"] = selected
        st.caption("This list is read from the models currently enabled for your Pangram API key.")
    elif api_key and Pangram is not None:
        st.caption("No Pangram model list is available yet.")

    st.divider()
    st.caption(
        "Pangram 4 accepts prose samples of at least 50 words. This app enforces that minimum and "
        "builds windows on sentence boundaries."
    )

    st.divider()
    st.header("Anthropic connection")
    anthropic_key = get_anthropic_api_key()
    if Anthropic is None:
        st.error("anthropic is not installed. Add it to requirements.txt and reboot the Streamlit app.")
    elif not anthropic_key:
        st.error("ANTHROPIC_API_KEY is not set in Streamlit Secrets.")
        st.caption("Add this on a new line in the same Secrets box:")
        st.code('ANTHROPIC_API_KEY = "paste-your-key-here"', language="toml")
        st.caption("Do not put the key in GitHub.")
    else:
        if st.session_state.get("anthropic_client") is None:
            try:
                with st.spinner("Connecting to Anthropic…"):
                    aclient, amodels = connect_anthropic(anthropic_key)
                st.session_state["anthropic_client"] = aclient
                st.session_state["anthropic_models"] = amodels
                st.session_state["anthropic_model"] = preferred_anthropic_model(amodels)
            except Exception as exc:
                st.session_state.pop("anthropic_client", None)
                st.session_state.pop("anthropic_models", None)
                st.session_state.pop("anthropic_model", None)
                st.error(f"Anthropic connection failed: {exc}")

        if st.session_state.get("anthropic_client") is not None:
            st.success("Anthropic API key loaded from Streamlit Secrets.")
            if st.button("Refresh Claude models", use_container_width=True):
                try:
                    with st.spinner("Refreshing Claude models…"):
                        aclient, amodels = connect_anthropic(anthropic_key)
                    st.session_state["anthropic_client"] = aclient
                    st.session_state["anthropic_models"] = amodels
                    current = st.session_state.get("anthropic_model")
                    if current not in amodels:
                        st.session_state["anthropic_model"] = preferred_anthropic_model(amodels)
                    st.rerun()
                except Exception as exc:
                    st.error(f"Could not refresh Claude models: {exc}")

    amodels = st.session_state.get("anthropic_models", [])
    if amodels:
        st.caption("Claude model is selected in Tab 1. The list is read live from Anthropic's Models API.")

claude_tab, experiment_tab, cal_tab, quick_tab, ab_tab, history_tab = st.tabs(
    [
        "1 · Claude → Pangram",
        "2 · Experiment Lab",
        "3 · Corpus calibration",
        "4 · 150-word microscope",
        "5 · Legacy A/B",
        "6 · History",
    ]
)


# -------------------------
# Tab 1: Claude → Pangram
# -------------------------
with claude_tab:
    st.subheader("One-click rapid fiction test")
    st.write(
        "Upload the drafting prompt and chapter packet, choose the API harness, and run. Standard mode preserves the "
        "old wrapper. Parity modes keep the two source documents separate and remove nearly all Streamlit-authored "
        "prompting before the prose is sent to Pangram. Prior outputs and Pangram scores are never included in generation."
    )
    st.info(
        "Default detector pass: sentence-aligned 150-word windows, no overlap, with undersized tail fragments excluded. "
        "For small jobs these go through Pangram realtime rather than the slower bulk queue."
    )

    g1, g2 = st.columns(2)
    with g1:
        prompt_upload = st.file_uploader(
            "Drafting prompt / method",
            type=["docx", "txt", "md"],
            accept_multiple_files=False,
            key="claude_prompt_upload_v23",
            help="Example: the current 6J drafting prompt.",
        )
    with g2:
        packet_upload = st.file_uploader(
            "Rapid-test chapter packet",
            type=["docx", "txt", "md"],
            accept_multiple_files=False,
            key="claude_packet_upload_v23",
            help="Example: Broken_Pump_Rapid_Test_Chapter_Packet_v22_1.docx.",
        )

    prompt_text = ""
    packet_text = ""
    prompt_name = "pasted prompt"
    packet_name = "pasted packet"
    if prompt_upload is not None:
        try:
            prompt_text = extract_text_from_upload(prompt_upload)
            prompt_name = prompt_upload.name
        except Exception as exc:
            st.error(f"Could not read drafting prompt: {exc}")
    if packet_upload is not None:
        try:
            packet_text = extract_text_from_upload(packet_upload)
            packet_name = packet_upload.name
        except Exception as exc:
            st.error(f"Could not read chapter packet: {exc}")

    with st.expander("Paste text instead of uploading files", expanded=False):
        if prompt_upload is None:
            prompt_text = st.text_area("Drafting prompt text", height=250, key="claude_prompt_paste_v23")
        if packet_upload is None:
            packet_text = st.text_area("Chapter packet text", height=250, key="claude_packet_paste_v23")

    detected_target = detect_word_count_target(packet_text)
    if detected_target:
        st.caption(f"Packet word-count target detected: **{detected_target:,} words**.")
    target_default = int(detected_target or CLAUDE_DEFAULT_TARGET_WORDS)

    # v2.9: one authoritative Claude model selector.  Earlier versions had a
    # sidebar selector and a Tab 1 selector both writing to anthropic_model,
    # which could fight each other across Streamlit reruns after a completed run.
    tab_models = st.session_state.get("anthropic_models", [])

    st.markdown("**Claude generation settings**")
    harness_mode = st.radio(
        "API harness",
        CLAUDE_HARNESS_MODES,
        index=1,
        horizontal=False,
        key="claude_harness_mode_v211",
        help=(
            "Standard preserves the old XML-style wrapper. Parity A sends the prompt and packet as separate "
            "content blocks with almost no added instruction. Parity B does the same and also uses a supplied "
            "Claude.ai system prompt."
        ),
    )

    claude_ai_system_prompt = ""
    if harness_mode == CLAUDE_HARNESS_PARITY_SYSTEM:
        st.info(
            "Parity B requires the published Claude.ai system prompt for the model you are testing. "
            "Paste it below exactly as published. The app will send it through the API's top-level system field."
        )
        st.markdown(f"[Open Anthropic's published Claude.ai system prompts]({CLAUDE_AI_SYSTEM_PROMPT_DOCS_URL})")
        claude_ai_system_prompt = st.text_area(
            "Claude.ai system prompt",
            height=260,
            key="claude_ai_system_prompt_v211",
            placeholder="Paste the published Opus 5 Claude.ai system prompt here.",
        )

    g1, g2, g3 = st.columns([2.0, 1.0, 1.0])
    with g1:
        if tab_models:
            model_key = "claude_model_tab_v29"
            if st.session_state.get(model_key) not in tab_models:
                preferred = preferred_anthropic_model(tab_models) or tab_models[0]
                st.session_state[model_key] = preferred
            selected_model = st.selectbox(
                "Claude model",
                tab_models,
                key=model_key,
                help="Authoritative model choice for this run. Change it between runs; no app restart is needed.",
            )
        else:
            selected_model = None
            st.selectbox("Claude model", ["Connect Anthropic first"], disabled=True, key="claude_model_tab_disabled_v29")
    with g2:
        effort = st.selectbox(
            "Claude effort",
            CLAUDE_EFFORT_LEVELS,
            index=CLAUDE_EFFORT_LEVELS.index(CLAUDE_DEFAULT_EFFORT),
            key="claude_effort_v29",
            help="For current Claude 5 models, effort controls thinking depth. High is Anthropic's default.",
        )
    with g3:
        max_tokens = st.number_input(
            "Max output tokens",
            min_value=1024,
            max_value=16000,
            value=CLAUDE_DEFAULT_MAX_TOKENS,
            step=512,
            key="claude_max_tokens_v29",
        )

    r1, r2 = st.columns(2)
    with r1:
        run_label = st.text_input("Run label", value="rapid-test", key="claude_run_label_v29")
    with r2:
        target_words = st.number_input(
            "Claude target words",
            min_value=150,
            max_value=3000,
            value=target_default,
            step=50,
            key=f"claude_target_words_v29_{target_default}",
        )

    if st.button(
        "Test selected Claude model",
        key="claude_model_ping_v210",
        help="Sends a tiny no-Pangram request so you can verify model access before a long run.",
    ):
        ping_client = st.session_state.get("anthropic_client")
        if not ping_client or not selected_model:
            st.error("Anthropic is not connected yet.")
        else:
            try:
                ping_kwargs = {
                    "model": selected_model,
                    "max_tokens": 32,
                    "messages": [{"role": "user", "content": "Reply with exactly: OK"}],
                }
                if selected_model in {"claude-fable-5", "claude-opus-5", "claude-sonnet-5"}:
                    ping_kwargs["output_config"] = {"effort": str(effort)}
                ping = ping_client.messages.create(**ping_kwargs)
                ping_stop, ping_cat, ping_exp = claude_stop_details(ping)
                ping_text = "".join(
                    getattr(b, "text", "")
                    for b in getattr(ping, "content", [])
                    if getattr(b, "type", None) == "text"
                ).strip()
                if ping_stop == "refusal":
                    details = " · ".join(x for x in [ping_cat, ping_exp] if x) or "no details"
                    st.error(f"{selected_model} is reachable but refused the tiny test: {details}")
                else:
                    st.success(f"{selected_model} responded: {ping_text or '(no text)'} · stop_reason={ping_stop}")
            except Exception as exc:
                st.error(f"Model test failed for {selected_model}: {exc}")

    request_bundle: dict[str, Any] | None = None
    if prompt_text and packet_text:
        request_bundle = build_claude_request(
            prompt_text,
            packet_text,
            int(target_words),
            harness_mode,
            claude_ai_system_prompt,
        )
        with st.expander("Exact Claude request content", expanded=False):
            st.text_area(
                "Request preview",
                value=request_bundle["preview"],
                height=420,
                disabled=True,
                key="claude_payload_preview_v211",
            )
            st.caption(
                "No prior conversation turns, prior outputs, or Pangram feedback are appended. "
                "In Parity A/B, the target-word control is logged but is not injected into the prompt; "
                "the packet's own target governs length."
            )

    pangram_client, pangram_model = get_connected_client()
    anthropic_client = st.session_state.get("anthropic_client")
    anthropic_model = selected_model
    if anthropic_model:
        mode_note = f"adaptive/default thinking · effort {effort}" if anthropic_model in {"claude-fable-5", "claude-opus-5", "claude-sonnet-5"} else "model-default controls"
        st.info(f"Claude request: **{anthropic_model}** · **{mode_note}** · **{harness_mode}**")
        if harness_mode != CLAUDE_HARNESS_STANDARD and anthropic_model != "claude-opus-5":
            st.warning(
                "The successful web comparison was generated with Opus 5. You can still choose another model, "
                "but that changes an additional variable."
            )
    system_prompt_ready = harness_mode != CLAUDE_HARNESS_PARITY_SYSTEM or bool(claude_ai_system_prompt.strip())
    ready = bool(
        prompt_text.strip() and packet_text.strip() and request_bundle and system_prompt_ready and
        anthropic_client and anthropic_model and pangram_client and pangram_model
    )

    if harness_mode == CLAUDE_HARNESS_PARITY_SYSTEM and not claude_ai_system_prompt.strip():
        st.warning("Paste the published Claude.ai system prompt above before running Parity B.")
    if not anthropic_client or not anthropic_model:
        st.warning("Connect Anthropic in the sidebar before running the test.")
    if not pangram_client or not pangram_model:
        st.warning("Connect Pangram in the sidebar before running the test.")

    st.caption("Change the model, files, effort, word target, or token limit at any time and click Run again. No restart is required.")

    if st.button(
        "Generate with Claude and score with Pangram",
        type="primary",
        disabled=not ready,
        key="claude_generate_score_v29",
    ):
        # Clear the prior displayed result immediately. A failed new run must never
        # leave an older run masquerading as the latest result.
        st.session_state["claude_pangram_last_v23"] = None
        run_id = str(uuid.uuid4())
        live = st.empty()
        run_status = st.status(f"Claude generation running…  Prompt: {prompt_name}", state="running")
        try:
            output_text, final_message, generation_seconds, generation_mode = run_claude_stream(
                anthropic_client,
                anthropic_model,
                request_bundle["messages"],
                max_tokens=int(max_tokens),
                effort=effort,
                live_placeholder=live,
                system_prompt=request_bundle.get("system"),
            )

            stop_reason, refusal_category, refusal_explanation = claude_stop_details(final_message)
            usage = getattr(final_message, "usage", None)
            input_tokens = getattr(usage, "input_tokens", None) if usage else None
            output_tokens = getattr(usage, "output_tokens", None) if usage else None

            # Fable 5 (and some other current models) can return a successful HTTP 200
            # with stop_reason="refusal" and no prose. Surface that explicitly instead
            # of misreporting it as a mysterious zero-word generation failure.
            if stop_reason == "refusal":
                detail_bits = []
                if refusal_category:
                    detail_bits.append(f"category: {refusal_category}")
                if refusal_explanation:
                    detail_bits.append(str(refusal_explanation))
                detail_text = " · ".join(detail_bits) if detail_bits else "No refusal details were supplied by the API."
                run_status.update(label=f"{anthropic_model} refused this request.", state="error")
                st.error(f"Claude refusal from {anthropic_model}: {detail_text}")
                st.caption(
                    f"Requested model: {anthropic_model} · {generation_mode} · "
                    f"input tokens: {input_tokens or 0:,} · output tokens: {output_tokens or 0:,}. "
                    "No automatic fallback was used, so this remains a clean model test."
                )
                st.session_state["claude_pangram_last_v23"] = None
                st.stop()

            output_wc = count_words(output_text)
            if output_wc < MIN_PANGRAM_WORDS:
                raise RuntimeError(
                    f"{anthropic_model} ended with stop_reason={stop_reason!r} and returned only {output_wc} words, "
                    f"below Pangram's {MIN_PANGRAM_WORDS}-word minimum."
                )

            generated_doc = SourceDoc(
                name=f"{run_label or 'rapid-test'}_claude.txt",
                expected_label="Claude API",
                text=output_text,
            )
            score_windows, score_tails = split_experiment_windows(
                [generated_doc],
                CLAUDE_PANGRAM_WINDOW_WORDS,
                0,
                CLAUDE_MAX_SCORE_WINDOWS,
                CLAUDE_PANGRAM_MIN_RATIO,
            )
            if not score_windows:
                score_windows = build_whole_document_windows([generated_doc])

            run_status.update(label="Claude complete. Pangram scoring running…", state="running")
            p_started = time.perf_counter()
            successes, failures = run_realtime_scan(pangram_client, pangram_model, score_windows)
            pangram_seconds = time.perf_counter() - p_started
            if not successes:
                raise RuntimeError("Pangram returned no successful score windows. " + "; ".join(str(x.get("error")) for x in failures))

            save_results(
                successes,
                experiment_name=run_label or "Claude rapid test",
                mode="Claude API auto-score",
                model=pangram_model,
                experiment_id=run_id,
            )
            summary = generation_score_summary(successes)
            record = {
                "run_id": run_id,
                "run_at": datetime.now(timezone.utc).isoformat(),
                "label": run_label,
                "anthropic_model": anthropic_model,
                "thinking_mode": f"{generation_mode} · {harness_mode}",
                "harness_mode": harness_mode,
                "request_description": request_bundle.get("description") if request_bundle else "",
                "system_prompt_supplied": bool(request_bundle and request_bundle.get("system")),
                "system_prompt_text": request_bundle.get("system") if request_bundle else None,
                "target_words": int(target_words),
                "output_words": output_wc,
                "max_tokens": int(max_tokens),
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "stop_reason": str(stop_reason or ""),
                "generation_seconds": generation_seconds,
                "pangram_model": pangram_model,
                "pangram_seconds": pangram_seconds,
                "pangram_windows": len(successes),
                "mean_ai": summary["mean_ai"],
                "mean_ai_involvement": summary["mean_involvement"],
                "prompt_name": prompt_name,
                "packet_name": packet_name,
                "prompt_text": prompt_text,
                "packet_text": packet_text,
                "output_text": output_text,
            }
            save_generation_run(record)
            st.session_state["claude_pangram_last_v23"] = {
                "record": record,
                "output_text": output_text,
                "rows": successes,
                "failures": failures,
                "tails": score_tails,
                "summary": summary,
            }
            run_status.update(label="Claude generation and Pangram scoring complete.", state="complete")
        except Exception as exc:
            run_status.update(label="Rapid test failed.", state="error")
            st.error(f"Rapid test failed: {exc}")

    last = st.session_state.get("claude_pangram_last_v23")
    if last:
        rec = last["record"]
        summary = last["summary"]
        st.markdown("### Latest rapid-test result")
        r1, r2, r3, r4 = st.columns(4)
        r1.metric("Claude generation", f"{rec['generation_seconds']:.1f} sec")
        r2.metric("Output", f"{rec['output_words']:,} words")
        r3.metric("Pangram scan", f"{rec['pangram_seconds']:.1f} sec")
        r4.metric(
            "Weighted AI fraction",
            f"{100*summary['mean_ai']:.1f}%" if summary["mean_ai"] is not None else "n/a",
        )

        st.caption(
            f"Prompt: {rec.get('prompt_name') or 'n/a'} · Packet: {rec.get('packet_name') or 'n/a'} · "
            f"Run ID: {rec.get('run_id') or 'n/a'}"
        )
        st.caption(
            f"Claude: {rec['anthropic_model']} · {rec['thinking_mode']} · "
            f"Pangram: {rec['pangram_model']} · {rec['pangram_windows']} scored 150-word window(s)."
        )
        if rec.get("input_tokens") is not None or rec.get("output_tokens") is not None:
            st.caption(
                f"Anthropic usage: {rec.get('input_tokens') or 0:,} input tokens · "
                f"{rec.get('output_tokens') or 0:,} output tokens · stop reason: {rec.get('stop_reason') or 'n/a'}."
            )
        st.write(
            f"Pangram windows: **{summary['human_windows']} Human · {summary['mixed_windows']} Mixed · "
            f"{summary['ai_windows']} AI**."
        )

        if st.button("Prepare next run", key=f"claude_prepare_next_v29_{rec.get('run_id', 'unknown')}"):
            st.session_state["claude_pangram_last_v23"] = None
            st.rerun()

        with st.expander("Generated prose", expanded=True):
            # Unique key per run prevents Streamlit from retaining text from an older run.
            st.text_area(
                "Claude output",
                value=last["output_text"],
                height=420,
                disabled=True,
                key=f"claude_output_v24_{rec.get('run_id', 'unknown')}",
            )
        prompt_stem = re.sub(r"[^A-Za-z0-9_-]+", "_", Path(rec.get("prompt_name") or "prompt").stem)[:60]
        run_short = str(rec.get("run_id") or "run")[:8]
        st.download_button(
            "Download Claude output as DOCX",
            data=prose_docx_bytes(last["output_text"]),
            file_name=f"{prompt_stem}_{run_short}_Claude.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"claude_docx_download_v24_{rec.get('run_id', 'unknown')}",
        )

        st.markdown("**Pangram 150-word windows**")
        show_result_table(results_dataframe(last["rows"]), key="claude_pangram_table_v23")
        if last.get("failures"):
            st.warning(f"{len(last['failures'])} Pangram window(s) failed.")
        if last.get("tails"):
            st.caption(f"Excluded {len(last['tails'])} undersized tail fragment(s) from the 150-word score.")

        result_bundle = {
            "run": rec,
            "pangram_summary": summary,
            "pangram_windows": [
                {k: v for k, v in row.items() if k != "raw_result"}
                for row in last["rows"]
            ],
        }
        st.download_button(
            "Download run record (JSON)",
            data=json.dumps(result_bundle, indent=2, ensure_ascii=False, default=str),
            file_name=f"{prompt_stem}_{run_short}_run.json",
            mime="application/json",
            key=f"claude_json_download_v24_{rec.get('run_id', 'unknown')}",
        )

    recent_generations = load_generation_runs(limit=20)
    if not recent_generations.empty:
        with st.expander("Recent Claude → Pangram runs", expanded=False):
            display_cols = [
                "run_at", "prompt_name", "label", "anthropic_model", "thinking_mode", "target_words", "output_words",
                "generation_seconds", "pangram_seconds", "pangram_windows", "mean_ai", "mean_ai_involvement",
            ]
            display_cols = [c for c in display_cols if c in recent_generations.columns]
            st.dataframe(
                recent_generations[display_cols].style.format({
                    "generation_seconds": "{:.1f}",
                    "pangram_seconds": "{:.1f}",
                    "mean_ai": "{:.1%}",
                    "mean_ai_involvement": "{:.1%}",
                }),
                use_container_width=True,
                hide_index=True,
            )
            st.download_button(
                "Download Claude → Pangram history CSV",
                data=recent_generations.to_csv(index=False).encode("utf-8"),
                file_name="claude_pangram_history.csv",
                mime="text/csv",
                key="claude_history_csv_v23",
            )


# -------------------------
# Tab 2: Experiment Lab
# -------------------------
with experiment_tab:
    st.subheader("Candidate-only prompt experiment")
    st.write(
        "v2.1 stops paying to rescan the parent on every experiment. Keep a reusable parent baseline, "
        "scan only the new candidate at ~500 words, and buy a whole-document scan only when the candidate merits it."
    )
    st.info(
        "Fast screen: candidate only, using score-eligible windows. Final check: candidate only, one whole-document "
        "request per file. Windows below 90% of the requested size are excluded from both the main score and the "
        "default API submission."
    )

    # ---------- Parent baseline ----------
    baselines = load_baselines()
    saved_versions = baselines["version"].astype(str).tolist() if not baselines.empty else []
    baseline_choice = st.selectbox(
        "Parent baseline source",
        ["Manual / current 6E baseline"] + saved_versions,
        key="lab_baseline_choice",
        help="Once a candidate is promoted, save it as a baseline. The next experiment can then avoid rescanning it.",
    )

    if baseline_choice != "Manual / current 6E baseline":
        brow = baselines[baselines["version"] == baseline_choice].iloc[0]
        parent_version = str(brow["version"])
        parent_screen_ai = float(brow["screen_ai"]) if pd.notna(brow["screen_ai"]) else None
        parent_screen_involvement = (
            float(brow["screen_ai_involvement"]) if pd.notna(brow["screen_ai_involvement"]) else None
        )
        parent_whole_ai = float(brow["whole_ai"]) if pd.notna(brow["whole_ai"]) else None
        st.caption(
            f"Loaded {parent_version}: "
            f"fast screen {100*parent_screen_ai:.1f}% AI"
            if parent_screen_ai is not None
            else f"Loaded {parent_version}: no fast-screen value saved."
        )
        if parent_whole_ai is not None:
            st.caption(f"Whole-document baseline: {100*parent_whole_ai:.1f}% AI.")
    else:
        b1, b2, b3 = st.columns(3)
        with b1:
            parent_version = st.text_input(
                "Parent version",
                value=BOOTSTRAP_PARENT_VERSION,
                key="lab_parent_version_manual",
            )
        with b2:
            parent_screen_pct = st.number_input(
                "Parent fast-screen AI %",
                min_value=0.0,
                max_value=100.0,
                value=100 * BOOTSTRAP_PARENT_SCREEN_AI,
                step=0.1,
                key="lab_parent_screen_pct",
                help="For the current 6E baseline this is the 500-word score after excluding <90%-size tails.",
            )
            parent_screen_ai = parent_screen_pct / 100.0
            parent_screen_involvement = None
        with b3:
            parent_whole_pct = st.number_input(
                "Parent whole-document AI %",
                min_value=0.0,
                max_value=100.0,
                value=100 * BOOTSTRAP_PARENT_WHOLE_AI,
                step=0.1,
                key="lab_parent_whole_pct",
            )
            parent_whole_ai = parent_whole_pct / 100.0

        if st.button("Save these parent values as baseline", key="lab_save_manual_baseline"):
            save_baseline(
                parent_version,
                screen_ai=parent_screen_ai,
                screen_ai_involvement=parent_screen_involvement,
                whole_ai=parent_whole_ai,
                source_note="Manual baseline entered in Experiment Lab",
            )
            st.success(f"Saved {parent_version} as a reusable baseline.")

    # ---------- Candidate metadata ----------
    cmeta1, cmeta2 = st.columns(2)
    with cmeta1:
        candidate_version = st.text_input(
            "Candidate version",
            value=BOOTSTRAP_CANDIDATE_VERSION,
            key="lab_candidate_version",
        )
    with cmeta2:
        test_set_note = st.text_input(
            "Test-set note",
            placeholder="Example: Ch. 3 development test; later use 3+ chapters/donors for promotion",
            key="lab_test_set_note",
        )

    change_note = st.text_area(
        "What changed in the candidate prompt?",
        placeholder="One controlled change only, if possible.",
        height=90,
        key="lab_change_note",
    )

    with st.expander("Store prompts with this experiment", expanded=False):
        p1, p2 = st.columns(2)
        with p1:
            parent_prompt = st.text_area("Parent prompt", height=260, key="lab_parent_prompt")
        with p2:
            candidate_prompt = st.text_area("Candidate prompt", height=260, key="lab_candidate_prompt")
        diff_text = prompt_diff(parent_prompt, candidate_prompt)
        if diff_text:
            st.markdown("**Prompt diff**")
            st.code(diff_text, language="diff")

    candidate_uploads = st.file_uploader(
        f"{candidate_version} output(s)",
        type=["docx", "txt", "md"],
        accept_multiple_files=True,
        key="lab_candidate_uploads_v21",
        help="One file is enough for development. Use 3+ different chapter/donor outputs before promoting a general method.",
    )
    candidate_docs, candidate_errors = source_docs_from_uploads(candidate_uploads, candidate_version)
    for err in candidate_errors:
        st.warning(err)

    # ---------- Screen settings ----------
    s1, s2, s3, s4 = st.columns(4)
    with s1:
        lab_size = st.selectbox(
            "Pangram window",
            ALL_SAMPLE_SIZES,
            index=ALL_SAMPLE_SIZES.index(EXPERIMENT_WINDOW_WORDS),
            key="lab_size_v21",
        )
    with s2:
        lab_overlap = st.slider(
            "Overlap",
            0,
            75,
            EXPERIMENT_OVERLAP_PCT,
            25,
            key="lab_overlap_v21",
        )
    with s3:
        lab_cap = st.number_input(
            "Max score windows / file",
            min_value=1,
            max_value=100,
            value=EXPERIMENT_MAX_WINDOWS_PER_FILE,
            step=1,
            key="lab_cap_v21",
        )
    with s4:
        structure_limit = st.slider(
            "Structure warning line",
            0.60,
            0.95,
            DEFAULT_STRUCTURE_SIMILARITY_LIMIT,
            0.01,
            key="lab_structure_limit_v21",
            help="Experimental heuristic; not a literary-quality score.",
        )

    candidate_main, candidate_tails = split_experiment_windows(
        candidate_docs,
        int(lab_size),
        int(lab_overlap),
        int(lab_cap),
        EXPERIMENT_MIN_WINDOW_RATIO,
    )
    whole_windows = build_whole_document_windows(candidate_docs)

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Candidate files", len(candidate_docs))
    m2.metric("Score windows", len(candidate_main))
    m3.metric("Excluded tails", len(candidate_tails))
    m4.metric("Fast-screen words", f"{sum(w.actual_words for w in candidate_main):,}")

    if candidate_main:
        st.markdown("**Fast-screen cost**")
        show_cost_estimate(candidate_main, key="lab_candidate_screen_cost")

    if whole_windows:
        whole_cost = estimate_bulk_cost(whole_windows)
        st.caption(
            f"Optional whole-document final: about **${whole_cost['estimated_cost']:.2f}** "
            f"for {whole_cost['actual_words']:,} word(s) across {whole_cost['requests']} file(s)."
        )

    tail_df = tail_window_table(candidate_tails, int(lab_size))
    if not tail_df.empty:
        with st.expander("Excluded tail fragments — not sent to Pangram", expanded=False):
            st.caption(
                f"These windows are under {EXPERIMENT_MIN_WINDOW_RATIO:.0%} of the requested size. "
                "They are shown only so you can see what was excluded; they cost $0."
            )
            st.dataframe(
                tail_df.style.format({"Target %": "{:.1%}"}),
                use_container_width=True,
                hide_index=True,
            )

    # Structural diversity is free because it is calculated locally.
    pre_diversity = candidate_diversity_table(candidate_docs)
    if len(candidate_docs) >= 2:
        with st.expander("Structural diversity check — no API cost", expanded=False):
            st.dataframe(
                pre_diversity.style.format({
                    "Paragraph-shape similarity": "{:.1%}",
                    "Sentence-shape similarity": "{:.1%}",
                    "Combined structural similarity": "{:.1%}",
                }),
                use_container_width=True,
                hide_index=True,
            )
            max_pre = float(pre_diversity["Combined structural similarity"].max())
            if max_pre >= structure_limit:
                st.warning(
                    f"At least one candidate pair is {max_pre:.1%} structurally similar. "
                    "Do not promote on Pangram score alone."
                )
            else:
                st.success(f"No candidate pair crosses the current {structure_limit:.0%} warning line.")
    elif len(candidate_docs) == 1:
        st.caption("One candidate can test detector behavior; cross-chapter structural repetition is not testable yet.")

    client, model = get_connected_client()

    # ---------- Stage 1: candidate-only fast screen ----------
    if st.button(
        f"1 · Run {candidate_version} fast screen",
        type="primary",
        disabled=not (client and model and candidate_main),
        key="lab_run_screen_v21",
    ):
        try:
            successes, failures = run_bulk_scan(client, model, candidate_main)
            run_id = str(uuid.uuid4())
            save_results(
                successes,
                experiment_name=f"{parent_version} → {candidate_version}",
                mode="Experiment Lab fast screen",
                model=model,
                experiment_id=run_id,
            )
            df = results_dataframe(successes)
            file_sum = experiment_file_summary(df)
            side_sum = side_summary(file_sum)
            crow = side_sum[side_sum["Version"] == candidate_version]
            candidate_mean = float(crow.iloc[0]["Mean AI fraction"]) if not crow.empty else None
            candidate_involvement = float(crow.iloc[0]["Mean AI involvement"]) if not crow.empty else None
            candidate_worst = float(crow.iloc[0]["Worst-file AI fraction"]) if not crow.empty else None
            delta_ai = (
                candidate_mean - parent_screen_ai
                if candidate_mean is not None and parent_screen_ai is not None
                else None
            )
            diversity_df = candidate_diversity_table(candidate_docs)
            max_sim = float(diversity_df["Combined structural similarity"].max()) if not diversity_df.empty else None
            verdict = experiment_verdict(delta_ai)

            record = {
                "experiment_id": run_id,
                "run_at": datetime.now(timezone.utc).isoformat(),
                "parent_version": parent_version,
                "candidate_version": candidate_version,
                "change_note": change_note,
                "test_set_note": test_set_note,
                "parent_prompt": parent_prompt,
                "candidate_prompt": candidate_prompt,
                "model": model,
                "target_words": int(lab_size),
                "overlap_pct": int(lab_overlap),
                "max_windows_per_file": int(lab_cap),
                "parent_files": 0,
                "candidate_files": len(candidate_docs),
                "parent_mean_ai": parent_screen_ai,
                "candidate_mean_ai": candidate_mean,
                "delta_ai": delta_ai,
                "candidate_worst_ai": candidate_worst,
                "candidate_max_structure_similarity": max_sim,
                "structure_similarity_limit": float(structure_limit),
                "verdict": verdict,
                "parent_whole_ai": parent_whole_ai,
                "candidate_whole_ai": None,
                "whole_delta_ai": None,
                "candidate_main_windows": len(candidate_main),
                "candidate_tail_windows": len(candidate_tails),
                "min_window_ratio": EXPERIMENT_MIN_WINDOW_RATIO,
            }
            save_experiment_run(record)
            st.session_state["lab_screen_last"] = {
                "run_id": run_id,
                "rows": successes,
                "failures": failures,
                "file_summary": file_sum,
                "side_summary": side_sum,
                "candidate_mean": candidate_mean,
                "candidate_involvement": candidate_involvement,
                "candidate_worst": candidate_worst,
                "delta_ai": delta_ai,
                "verdict": verdict,
                "parent_version": parent_version,
                "candidate_version": candidate_version,
                "parent_screen_ai": parent_screen_ai,
                "parent_whole_ai": parent_whole_ai,
                "change_note": change_note,
                "test_set_note": test_set_note,
                "parent_prompt": parent_prompt,
                "candidate_prompt": candidate_prompt,
                "diversity": diversity_df,
                "structure_limit": float(structure_limit),
                "record": record,
            }
        except Exception as exc:
            st.error(f"Fast screen failed: {exc}")

    screen_last = st.session_state.get("lab_screen_last")
    if screen_last and screen_last.get("candidate_version") == candidate_version:
        st.divider()
        st.subheader("Fast-screen result")
        candidate_mean = screen_last["candidate_mean"]
        delta_ai = screen_last["delta_ai"]
        r1, r2, r3 = st.columns(3)
        r1.metric(
            f"{screen_last['parent_version']} baseline",
            f"{100*screen_last['parent_screen_ai']:.1f}% AI"
            if screen_last["parent_screen_ai"] is not None
            else "not set",
        )
        r2.metric(
            f"{candidate_version} fast screen",
            f"{100*candidate_mean:.1f}% AI" if candidate_mean is not None else "n/a",
        )
        r3.metric(
            "Candidate change",
            f"{100*delta_ai:+.1f} points" if delta_ai is not None else "n/a",
            delta_color="inverse",
        )
        verdict = screen_last["verdict"]
        if verdict in {"STRONG IMPROVEMENT", "PROMISING"}:
            st.success(f"Fast-screen result: **{verdict}**")
        elif verdict == "INCONCLUSIVE":
            st.info("Fast-screen result: **INCONCLUSIVE** — close candidates require the whole-document check.")
        else:
            st.error(f"Fast-screen result: **{verdict}**")

        st.dataframe(
            screen_last["file_summary"].style.format({
                "Weighted AI fraction": "{:.1%}",
                "Weighted AI involvement": "{:.1%}",
                "Human %": "{:.1f}",
                "Mixed %": "{:.1f}",
                "AI %": "{:.1f}",
            }),
            use_container_width=True,
            hide_index=True,
        )

    # ---------- Stage 2: candidate-only whole-document final ----------
    st.markdown("### Optional final check")
    st.caption(
        "Run this only when the candidate is worth a full-document decision. The parent is not rescanned."
    )
    if whole_windows:
        show_cost_estimate(whole_windows, key="lab_whole_cost")

    if st.button(
        f"2 · Run {candidate_version} whole-document final",
        disabled=not (client and model and whole_windows),
        key="lab_run_whole_v21",
    ):
        try:
            whole_successes, whole_failures = run_bulk_scan(client, model, whole_windows)
            screen_last = st.session_state.get("lab_screen_last")
            if screen_last and screen_last.get("candidate_version") == candidate_version:
                run_id = screen_last["run_id"]
                record = dict(screen_last["record"])
                candidate_screen_ai = screen_last["candidate_mean"]
                candidate_screen_involvement = screen_last["candidate_involvement"]
            else:
                run_id = str(uuid.uuid4())
                record = {
                    "experiment_id": run_id,
                    "run_at": datetime.now(timezone.utc).isoformat(),
                    "parent_version": parent_version,
                    "candidate_version": candidate_version,
                    "change_note": change_note,
                    "test_set_note": test_set_note,
                    "parent_prompt": parent_prompt,
                    "candidate_prompt": candidate_prompt,
                    "model": model,
                    "target_words": int(lab_size),
                    "overlap_pct": int(lab_overlap),
                    "max_windows_per_file": int(lab_cap),
                    "parent_files": 0,
                    "candidate_files": len(candidate_docs),
                    "parent_mean_ai": parent_screen_ai,
                    "candidate_mean_ai": None,
                    "delta_ai": None,
                    "candidate_worst_ai": None,
                    "candidate_max_structure_similarity": (
                        float(pre_diversity["Combined structural similarity"].max())
                        if not pre_diversity.empty else None
                    ),
                    "structure_similarity_limit": float(structure_limit),
                    "verdict": "WHOLE-DOCUMENT ONLY",
                    "candidate_main_windows": len(candidate_main),
                    "candidate_tail_windows": len(candidate_tails),
                    "min_window_ratio": EXPERIMENT_MIN_WINDOW_RATIO,
                }
                candidate_screen_ai = None
                candidate_screen_involvement = None

            save_results(
                whole_successes,
                experiment_name=f"{parent_version} → {candidate_version}",
                mode="Experiment Lab whole document",
                model=model,
                experiment_id=run_id,
            )
            whole_df = whole_document_summary(results_dataframe(whole_successes))
            candidate_whole_ai = mean_whole_ai(whole_df)
            whole_delta_ai = (
                candidate_whole_ai - parent_whole_ai
                if candidate_whole_ai is not None and parent_whole_ai is not None
                else None
            )
            record.update(
                {
                    "parent_whole_ai": parent_whole_ai,
                    "candidate_whole_ai": candidate_whole_ai,
                    "whole_delta_ai": whole_delta_ai,
                }
            )
            save_experiment_run(record)

            st.session_state["lab_whole_last"] = {
                "run_id": run_id,
                "rows": whole_successes,
                "failures": whole_failures,
                "whole_summary": whole_df,
                "candidate_whole_ai": candidate_whole_ai,
                "whole_delta_ai": whole_delta_ai,
                "candidate_screen_ai": candidate_screen_ai,
                "candidate_screen_involvement": candidate_screen_involvement,
                "parent_version": parent_version,
                "candidate_version": candidate_version,
                "parent_whole_ai": parent_whole_ai,
                "candidate_prompt": candidate_prompt,
                "test_set_note": test_set_note,
            }
        except Exception as exc:
            st.error(f"Whole-document scan failed: {exc}")

    whole_last = st.session_state.get("lab_whole_last")
    if whole_last and whole_last.get("candidate_version") == candidate_version:
        st.subheader("Whole-document result")
        w1, w2, w3 = st.columns(3)
        w1.metric(
            f"{whole_last['parent_version']} final baseline",
            f"{100*whole_last['parent_whole_ai']:.1f}% AI"
            if whole_last["parent_whole_ai"] is not None
            else "not set",
        )
        w2.metric(
            f"{candidate_version} whole document",
            f"{100*whole_last['candidate_whole_ai']:.1f}% AI"
            if whole_last["candidate_whole_ai"] is not None
            else "n/a",
        )
        w3.metric(
            "Final change",
            f"{100*whole_last['whole_delta_ai']:+.1f} points"
            if whole_last["whole_delta_ai"] is not None
            else "n/a",
            delta_color="inverse",
        )
        st.dataframe(
            whole_last["whole_summary"].style.format({
                "AI fraction": "{:.1%}",
                "AI involvement": "{:.1%}",
                "Human fraction": "{:.1%}",
            }),
            use_container_width=True,
            hide_index=True,
        )

        if screen_last and screen_last.get("candidate_version") == candidate_version:
            st.caption(
                "Decision rule: when the fast screen and whole-document result disagree, use the whole-document result "
                "for the detector decision and keep the fast screen as a cheap directional test."
            )

        if st.button(
            f"Promote {candidate_version} as reusable parent baseline",
            key="lab_promote_baseline_v21",
        ):
            save_baseline(
                candidate_version,
                screen_ai=whole_last.get("candidate_screen_ai"),
                screen_ai_involvement=whole_last.get("candidate_screen_involvement"),
                whole_ai=whole_last.get("candidate_whole_ai"),
                source_note=whole_last.get("test_set_note") or f"Promoted from Experiment Lab",
                prompt_text=whole_last.get("candidate_prompt") or "",
            )
            st.success(
                f"Saved {candidate_version}. On the next experiment, select it from Parent baseline source; "
                "the app will not rescan it."
            )

    # ---------- Handoff / download ----------
    screen_last = st.session_state.get("lab_screen_last")
    whole_last = st.session_state.get("lab_whole_last")
    if (
        (screen_last and screen_last.get("candidate_version") == candidate_version)
        or (whole_last and whole_last.get("candidate_version") == candidate_version)
    ):
        lines = [
            f"Pangram Experiment Lab: {parent_version} → {candidate_version}",
            f"Change tested: {change_note or '(not entered)'}",
        ]
        if screen_last and screen_last.get("candidate_version") == candidate_version:
            lines.append(
                f"Fast screen: parent {100*parent_screen_ai:.1f}% AI → "
                f"candidate {100*screen_last['candidate_mean']:.1f}% AI "
                f"({100*screen_last['delta_ai']:+.1f} points)."
                if screen_last["candidate_mean"] is not None and screen_last["delta_ai"] is not None
                else "Fast screen completed; comparison unavailable."
            )
            lines.append(
                f"Score windows: {len(candidate_main)}; excluded tails: {len(candidate_tails)} "
                f"(<{EXPERIMENT_MIN_WINDOW_RATIO:.0%} target size)."
            )
        if whole_last and whole_last.get("candidate_version") == candidate_version:
            if whole_last["candidate_whole_ai"] is not None:
                if whole_last["whole_delta_ai"] is not None:
                    lines.append(
                        f"Whole document: parent {100*parent_whole_ai:.1f}% AI → "
                        f"candidate {100*whole_last['candidate_whole_ai']:.1f}% AI "
                        f"({100*whole_last['whole_delta_ai']:+.1f} points)."
                    )
                else:
                    lines.append(f"Whole document candidate: {100*whole_last['candidate_whole_ai']:.1f}% AI.")
        if not pre_diversity.empty:
            lines.append(
                f"Max candidate structural similarity: "
                f"{100*float(pre_diversity['Combined structural similarity'].max()):.1f}% "
                f"(warning line {100*structure_limit:.0f}%)."
            )
        else:
            lines.append("Structural diversity: not testable from fewer than 2 candidate files.")

        handoff = "\n".join(lines)
        st.text_area("Copy back to ChatGPT", value=handoff, height=220, key="lab_handoff_v21")

        export_rows = []
        if screen_last and screen_last.get("candidate_version") == candidate_version:
            for row in screen_last["rows"]:
                r = {k: v for k, v in row.items() if k != "raw_result"}
                r["lab_stage"] = "fast_screen"
                export_rows.append(r)
        if whole_last and whole_last.get("candidate_version") == candidate_version:
            for row in whole_last["rows"]:
                r = {k: v for k, v in row.items() if k != "raw_result"}
                r["lab_stage"] = "whole_document"
                export_rows.append(r)
        if export_rows:
            st.download_button(
                "Download experiment CSV",
                pd.DataFrame(export_rows).to_csv(index=False).encode("utf-8"),
                file_name=f"pangram_experiment_{parent_version}_to_{candidate_version}.csv",
                mime="text/csv",
                key="lab_download_v21",
            )

# -------------------------
# Tab 2: Calibration
# -------------------------
with cal_tab:
    st.subheader("Find the smallest useful test window")
    st.write(
        "Load writing you know is human and writing you know was generated by AI. The app creates "
        "mechanical, sentence-aligned windows at several sizes and sends them to Pangram in bulk."
    )

    c1, c2 = st.columns(2)
    with c1:
        human_files = st.file_uploader(
            "Known-human corpus",
            type=["docx", "txt", "md"],
            accept_multiple_files=True,
            key="human_uploads",
        )
    with c2:
        ai_files = st.file_uploader(
            "Known-AI corpus",
            type=["docx", "txt", "md"],
            accept_multiple_files=True,
            key="ai_uploads",
        )

    human_docs, human_errors = source_docs_from_uploads(human_files, "Human")
    ai_docs, ai_errors = source_docs_from_uploads(ai_files, "AI")
    for err in human_errors + ai_errors:
        st.warning(err)

    settings1, settings2, settings3 = st.columns(3)
    with settings1:
        sample_sizes = st.multiselect(
            "Target window sizes (words)",
            ALL_SAMPLE_SIZES,
            default=DEFAULT_SAMPLE_SIZES,
        )
    with settings2:
        overlap_pct = st.slider("Window overlap", 0, 75, DEFAULT_OVERLAP_PCT, 25)
    with settings3:
        cap_per = st.number_input(
            "Max windows / file / size",
            min_value=1,
            max_value=200,
            value=CALIBRATION_DEFAULT_MAX_WINDOWS,
            step=1,
            help="Windows are selected evenly across each file when this cap is reached.",
        )

    docs = human_docs + ai_docs
    windows = make_calibration_windows(
        docs,
        [s for s in sample_sizes if s >= MIN_PANGRAM_WORDS],
        overlap_pct / 100.0,
        int(cap_per),
    ) if docs and sample_sizes else []

    total_words = sum(w.actual_words for w in windows)
    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Human files", len(human_docs))
    m2.metric("AI files", len(ai_docs))
    m3.metric("Pangram samples", len(windows))
    m4.metric("Words submitted", f"{total_words:,}")

    if windows:
        preview = pd.DataFrame([asdict(w) for w in windows])
        with st.expander("Preview the mechanical sample plan"):
            st.dataframe(
                preview[["source_name", "expected_label", "target_words", "actual_words", "sentence_start", "sentence_end"]],
                use_container_width=True,
                hide_index=True,
            )
        show_cost_estimate(windows, key="cal_cost")

    experiment_name = st.text_input(
        "Experiment name",
        value=f"Calibration {datetime.now().strftime('%Y-%m-%d')}",
        key="cal_experiment_name",
    )

    client, model = get_connected_client()
    run_disabled = not (client and model and windows and human_docs and ai_docs)
    if st.button("Run calibration through Pangram", type="primary", disabled=run_disabled):
        try:
            successes, failures = run_bulk_scan(client, model, windows)
            exp_id = save_results(
                successes,
                experiment_name=experiment_name,
                mode="Calibration",
                model=model,
            )
            st.session_state["cal_last_rows"] = successes
            st.session_state["cal_last_failures"] = failures
            st.session_state["cal_last_experiment_id"] = exp_id
        except Exception as exc:
            st.error(f"Pangram run failed: {exc}")

    rows = st.session_state.get("cal_last_rows", [])
    failures = st.session_state.get("cal_last_failures", [])
    if rows:
        df = results_dataframe(rows)
        st.divider()
        st.subheader("Calibration result")

        summary = calibration_summary(df)
        sep = separation_summary(df)

        if not sep.empty:
            st.markdown("**Separation by test length**")
            st.dataframe(sep, use_container_width=True, hide_index=True)
            chart = sep.set_index("Target words")[["Human correctly Human %", "AI correctly AI %"]]
            st.line_chart(chart)

            tc1, tc2 = st.columns(2)
            with tc1:
                human_thresh = st.slider("Required Human-control success", 0.50, 1.00, 0.90, 0.05)
            with tc2:
                ai_thresh = st.slider("Required AI-control success", 0.50, 1.00, 0.90, 0.05)
            st.info(recommend_size(sep, human_thresh, ai_thresh))

        with st.expander("Detailed calibration summary"):
            st.dataframe(summary, use_container_width=True, hide_index=True)

        with st.expander("Every tested window"):
            show_result_table(df, "cal_results_df")

        st.download_button(
            "Download calibration CSV",
            data=df.to_csv(index=False).encode("utf-8"),
            file_name="pangram_calibration_results.csv",
            mime="text/csv",
        )

        if failures:
            st.warning(f"{len(failures)} Pangram item(s) failed.")
            st.dataframe(pd.DataFrame(failures), use_container_width=True, hide_index=True)

    elif not run_disabled and not rows:
        st.caption("Ready. Run the calibration when you want to spend the API calls.")
    elif not client:
        st.info("Connect to Pangram in the sidebar first.")
    elif not (human_docs and ai_docs):
        st.info("Add at least one known-human and one known-AI document.")


# -------------------------
# Tab 3: 150-word microscope
# -------------------------
with quick_tab:
    st.subheader("150-word microscope")
    st.write(
        "Use this for local diagnosis, not for ranking prompt versions. The 6C/6D calibration showed that "
        "150-word windows can disagree with the whole-document direction."
    )

    quick_upload = st.file_uploader(
        "Optional DOCX / TXT / MD",
        type=["docx", "txt", "md"],
        accept_multiple_files=False,
        key="quick_upload",
    )
    default_quick_text = ""
    quick_name = "pasted_text.txt"
    if quick_upload is not None:
        try:
            default_quick_text = extract_text_from_upload(quick_upload)
            quick_name = quick_upload.name
        except Exception as exc:
            st.error(str(exc))

    quick_text = st.text_area(
        "Text to test",
        value=default_quick_text,
        height=260,
        key="quick_text",
    )

    q1, q2, q3 = st.columns(3)
    with q1:
        q_size = st.selectbox("Target words", ALL_SAMPLE_SIZES, index=ALL_SAMPLE_SIZES.index(CALIBRATED_WINDOW_WORDS))
    with q2:
        q_overlap = st.slider("Overlap", 0, 75, DEFAULT_OVERLAP_PCT, 25, key="quick_overlap")
    with q3:
        q_cap = st.number_input("Max windows", 1, 100, QUICK_DEFAULT_MAX_WINDOWS, 1, key="quick_cap")

    qdoc = SourceDoc(name=quick_name, expected_label="Unknown", text=clean_text(quick_text))
    qwindows = evenly_cap(
        build_sentence_windows(qdoc, int(q_size), q_overlap / 100.0),
        int(q_cap),
    ) if count_words(qdoc.text) >= MIN_PANGRAM_WORDS else []

    st.caption(
        f"{count_words(qdoc.text):,} source words → {len(qwindows)} Pangram window(s) → "
        f"{sum(w.actual_words for w in qwindows):,} submitted words."
    )
    show_cost_estimate(qwindows, key="quick_cost")

    qexp = st.text_input("Experiment name", value="Quick scan", key="quick_exp")
    client, model = get_connected_client()
    if st.button(
        "Scan these windows",
        type="primary",
        disabled=not (client and model and qwindows),
        key="quick_run",
    ):
        try:
            successes, failures = run_realtime_scan(client, model, qwindows)
            save_results(successes, experiment_name=qexp, mode="Quick", model=model)
            st.session_state["quick_last_rows"] = successes
            st.session_state["quick_last_failures"] = failures
        except Exception as exc:
            st.error(f"Pangram run failed: {exc}")

    qrows = st.session_state.get("quick_last_rows", [])
    if qrows:
        qdf = results_dataframe(qrows)
        qcounts = qdf["prediction"].value_counts().to_dict()
        cols = st.columns(3)
        cols[0].metric("Human windows", qcounts.get("Human", 0))
        cols[1].metric("Mixed windows", qcounts.get("Mixed", 0))
        cols[2].metric("AI windows", qcounts.get("AI", 0))
        show_result_table(qdf, "quick_results_df")
        st.download_button(
            "Download quick-scan CSV",
            qdf.to_csv(index=False).encode("utf-8"),
            file_name="pangram_quick_scan.csv",
            mime="text/csv",
        )


# -------------------------
# Tab 4: Legacy A/B prompt test
# -------------------------
with ab_tab:
    st.subheader("Legacy A/B prompt experiment")
    st.write(
        "Paste comparable output from the current prompt and one candidate prompt. The app uses the same "
        "window settings on both and compares their Pangram distributions."
    )

    a_col, b_col = st.columns(2)
    with a_col:
        control_text = st.text_area("Control output", height=300, key="control_text")
    with b_col:
        candidate_text = st.text_area("Candidate output", height=300, key="candidate_text")

    ab1, ab2, ab3 = st.columns(3)
    with ab1:
        ab_size = st.selectbox("Target words", ALL_SAMPLE_SIZES, index=ALL_SAMPLE_SIZES.index(CALIBRATED_WINDOW_WORDS), key="ab_size")
    with ab2:
        ab_overlap = st.slider("Overlap", 0, 75, DEFAULT_OVERLAP_PCT, 25, key="ab_overlap")
    with ab3:
        ab_cap = st.number_input("Max windows per side", 1, 100, QUICK_DEFAULT_MAX_WINDOWS, 1, key="ab_cap")

    cdoc = SourceDoc("Control", "Control", clean_text(control_text))
    ndoc = SourceDoc("Candidate", "Candidate", clean_text(candidate_text))
    cwindows = evenly_cap(build_sentence_windows(cdoc, int(ab_size), ab_overlap / 100.0), int(ab_cap)) if count_words(cdoc.text) >= MIN_PANGRAM_WORDS else []
    nwindows = evenly_cap(build_sentence_windows(ndoc, int(ab_size), ab_overlap / 100.0), int(ab_cap)) if count_words(ndoc.text) >= MIN_PANGRAM_WORDS else []
    abwindows = cwindows + nwindows
    show_realtime_cost_estimate(abwindows, key="ab_cost")
    if abwindows:
        st.caption("Legacy A/B runs these small paired tests in realtime; the rest of the app can still use bulk where it saves meaningful money.")

    abexp = st.text_input("Experiment name", value="Prompt A/B", key="ab_exp")
    client, model = get_connected_client()
    if st.button(
        "Run A/B Pangram test",
        type="primary",
        disabled=not (client and model and cwindows and nwindows),
        key="ab_run",
    ):
        try:
            successes, failures = run_realtime_scan(client, model, abwindows)
            save_results(successes, experiment_name=abexp, mode="A/B realtime", model=model)
            st.session_state["ab_last_rows"] = successes
            st.session_state["ab_last_failures"] = failures
        except Exception as exc:
            st.error(f"Pangram run failed: {exc}")

    abrows = st.session_state.get("ab_last_rows", [])
    if abrows:
        abdf = results_dataframe(abrows)
        compare = []
        for label, group in abdf.groupby("expected_label"):
            compare.append(
                {
                    "Version": label,
                    "Windows": len(group),
                    "Human %": 100 * (group["prediction"] == "Human").mean(),
                    "Mixed %": 100 * (group["prediction"] == "Mixed").mean(),
                    "AI %": 100 * (group["prediction"] == "AI").mean(),
                    "Mean human fraction": group["fraction_human"].mean(),
                    "Mean AI fraction": group["fraction_ai"].mean(),
                    "Mean AI involvement": group["mean_ai_involvement"].mean(),
                }
            )
        cmp_df = pd.DataFrame(compare)
        st.dataframe(cmp_df, use_container_width=True, hide_index=True)
        if not cmp_df.empty:
            st.bar_chart(cmp_df.set_index("Version")[["Human %", "Mixed %", "AI %"]])
        with st.expander("Every A/B window"):
            show_result_table(abdf, "ab_results_df")


# -------------------------
# Tab 5: History
# -------------------------
with history_tab:
    st.subheader("Experiment history")
    baseline_hist = load_baselines()
    if not baseline_hist.empty:
        st.markdown("**Reusable parent baselines**")
        base_show = baseline_hist.copy()
        for c in ["screen_ai", "screen_ai_involvement", "whole_ai"]:
            if c in base_show.columns:
                base_show[c] = pd.to_numeric(base_show[c], errors="coerce")
        st.dataframe(base_show, use_container_width=True, hide_index=True)
        st.caption("These saved baselines let the next experiment scan only the new candidate.")
        st.divider()

    exp_hist = load_experiment_history()
    if not exp_hist.empty:
        st.markdown("**Experiment Lab runs**")
        exp_show = exp_hist.copy()
        for c in ["parent_mean_ai", "candidate_mean_ai", "delta_ai", "candidate_worst_ai", "candidate_max_structure_similarity"]:
            if c in exp_show.columns:
                exp_show[c] = pd.to_numeric(exp_show[c], errors="coerce")
        st.dataframe(exp_show, use_container_width=True, hide_index=True)
        st.download_button(
            "Download Experiment Lab history CSV",
            exp_show.to_csv(index=False).encode("utf-8"),
            file_name="pangram_experiment_lab_history.csv",
            mime="text/csv",
            key="download_lab_history",
        )
        st.divider()
    st.caption(
        "Streamlit Cloud can recycle an app's local filesystem. Treat this built-in SQLite history as "
        "working history, not permanent storage; download CSVs you want to keep."
    )
    hist = load_history()
    if hist.empty:
        st.info("No saved Pangram runs yet.")
    else:
        e1, e2, e3 = st.columns(3)
        e1.metric("Saved windows", len(hist))
        e2.metric("Experiments", hist["experiment_id"].nunique())
        e3.metric("Latest run", str(hist.iloc[0]["run_at"])[:19].replace("T", " "))

        names = ["All"] + sorted([x for x in hist["experiment_name"].dropna().unique()])
        selected_name = st.selectbox("Experiment", names)
        shown = hist if selected_name == "All" else hist[hist["experiment_name"] == selected_name]

        display_cols = [
            "run_at",
            "experiment_name",
            "mode",
            "model",
            "source_name",
            "expected_label",
            "target_words",
            "actual_words",
            "prediction",
            "fraction_human",
            "fraction_ai_assisted",
            "fraction_ai",
            "mean_ai_involvement",
            "max_humanizer_score",
            "text",
        ]
        st.dataframe(
            shown[[c for c in display_cols if c in shown.columns]],
            use_container_width=True,
            hide_index=True,
        )
        st.download_button(
            "Download history CSV",
            shown.to_csv(index=False).encode("utf-8"),
            file_name="pangram_microscope_history.csv",
            mime="text/csv",
        )

st.divider()
st.caption(
    "Interpretation rule: this is a measurement lab, not a rewriting loop. Keep the drafting model blind to Pangram results. "
    "Optimize for generalization across different chapters/donors, not for one detector-perfect structural template."
)
